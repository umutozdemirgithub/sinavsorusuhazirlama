import streamlit as st
import pandas as pd
import sqlite3
import json
import random
import os
import urllib.request
import base64
import bcrypt
import time
import zipfile
import re  # Regex
from contextlib import contextmanager
from io import BytesIO
from datetime import datetime
import difflib

# --- AI Kütüphaneleri için Hata Yönetimi ---
import google.generativeai as genai
from google.generativeai import types
from google.api_core import exceptions as api_exceptions # API Hata Yönetimi

# --- Kütüphane Kontrolleri ve Importlar ---
try:
    from fpdf import FPDF
except ImportError:
    st.error("❌ FPDF kütüphanesi yüklü değil! Konsola: pip install fpdf")
    st.stop()

try:
    import plotly.express as px
    import plotly.graph_objects as go
except ImportError:
    st.warning("⚠️ Grafikler için 'plotly' gerekli: pip install plotly")

try:
    import xlsxwriter
except ImportError:
    st.warning("⚠️ Excel işlemleri için 'xlsxwriter' gerekli: pip install xlsxwriter")

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    st.warning("⚠️ Word çıktısı için 'python-docx' gerekli: pip install python-docx")

try:
    from streamlit_option_menu import option_menu
except ImportError:
    st.warning("⚠️ Daha şık menü için: pip install streamlit-option-menu")

try:
    import PyPDF2
except ImportError:
    st.warning("⚠️ PDF okuma için 'PyPDF2' gerekli: pip install PyPDF2")


# --- API Anahtarını Alma ve Yapılandırma ---
try:
    api_key = st.secrets.get("GOOGLE_API_KEY")
    if api_key:
        genai.configure(api_key=api_key)
except Exception:
    api_key = None

try:
    import requests
    from bs4 import BeautifulSoup
except ImportError:
    st.warning("⚠️ Web veri çekme için 'requests' ve 'beautifulsoup4' gerekli: pip install requests beautifulsoup4")
# ==============================================================================
# 1. AYARLAR VE TASARIM
# ==============================================================================
st.set_page_config(
    page_title="SSOP Pro v5.2: Enterprise",
    page_icon="🎓",
    layout="wide",
    initial_sidebar_state="expanded"
)
DB_FILE = "ssop_v5.sqlite"
FONT_FILENAME = "DejaVuSans.ttf"
FONT_URL = "https://github.com/dejavu-fonts/dejavu-fonts/raw/master/ttf/DejaVuSans.ttf"
OPTION_KEY_REGEX = r'^[A-Za-z]$' 

MENU_ROLES = {
    "Gösterge Paneli": ["Admin", "Öğretim Üyesi"],
    "Dersler": ["Admin", "Öğretim Üyesi"],
    "Soru Ekle": ["Admin", "Öğretim Üyesi"],
    "Soru Bankası": ["Admin", "Öğretim Üyesi"],
    "Sınav Oluştur": ["Admin", "Öğretim Üyesi"],
    "Arşiv": ["Admin", "Öğretim Üyesi"],
    "Yönetim": ["Admin"],
}

def local_css():
    st.markdown("""
    <style>

    /* --- GENEL UYGULAMA ARKA PLANI --- */
    .stApp {
        background-color: #eef2f7;
    }

    /* --- SIDEBAR KUTUSU --- */
    [data-testid="stSidebar"] {
        background: linear-gradient(180deg, #1f2c3a 0%, #253645 100%);
        padding-top: 30px;
        box-shadow: 4px 0 12px rgba(0,0,0,0.15);
    }

    /* --- PROFESYONEL KULLANICI KARTI --- */
    .user-card {
        background: rgba(255,255,255,0.08);
        backdrop-filter: blur(6px);
        padding: 25px;
        border-radius: 18px;
        text-align: center;
        margin-bottom: 25px;
        border: 1px solid rgba(255,255,255,0.12);
        transition: 0.3s ease;
    }
    .user-card:hover {
        transform: translateY(-3px);
        box-shadow: 0 6px 20px rgba(0,0,0,0.25);
    }

    /* --- PROFİL FOTOĞRAFI --- */
    .profile-pic {
        width: 90px;
        height: 90px;
        border-radius: 50%;
        border: 4px solid #3498db;
        object-fit: cover;
        margin-bottom: 12px;
        transition: 0.3s;
    }
    .profile-pic:hover {
        border-color: #5dade2;
        transform: scale(1.05);
    }

    /* --- PROFİL ADI / ROLÜ --- */
    .user-name {
        font-size: 1.3em;
        font-weight: 600;
        color: #ecf0f1;
        margin-bottom: 3px;
    }
    .user-role {
        font-size: 0.95em;
        color: #bdc3c7;
        margin-top: 0;
    }

    /* --- MENU ICONLARI --- */
    .st-emotion-cache-1m3botj,
    .st-emotion-cache-16idsys {
        color: #f1c40f !important;
    }

    /* --- MENU STILİ --- */
    .st-emotion-cache-1dp5vir {
        background-color: transparent !important;
    }

    /* --- ÇIKIŞ BUTONU --- */
    .exit-btn > button {
        background: #e74c3c;
        color: white;
        border: none;
        padding: 0.6em;
        border-radius: 10px;
        font-weight: 600;
        transition: 0.25s ease;
    }
    .exit-btn > button:hover {
        background: #c0392b;
        transform: translateY(-2px);
    }

    </style>
    """, unsafe_allow_html=True)

# ==============================================================================
# 2. VERİTABANI VE ALTYAPI
# ==============================================================================
@st.cache_resource
def check_and_download_font():
    """Font dosyasını sadece yoksa indirir ve cache mekanizmasını kullanır."""
    if not os.path.exists(FONT_FILENAME):
        try:
            urllib.request.urlretrieve(FONT_URL, FONT_FILENAME)
        except Exception:
            pass

def hash_password(password):
    return bcrypt.hashpw(password.encode(), bcrypt.gensalt()).decode()

def check_password(password, hashed):
    try:
        return bcrypt.checkpw(password.encode(), hashed.encode())
    except:
        return False

# ==============================================================================
# 3. YARDIMCI SINIFLAR (AI / PDF / WORD)
# ==============================================================================
# ==============================================================================
# AI CLASS (Geliştirme 1: Gelişmiş API Anahtarını Yönetimi)
# ==============================================================================
class AIGenerator:
    @staticmethod
    def get_api_key(provider):
        try:
            if provider == "google": 
                if f"user_provided_{provider}_key" in st.session_state and st.session_state[f"user_provided_{provider}_key"]:
                    return st.session_state[f"user_provided_{provider}_key"]
                return st.secrets.get("GOOGLE_API_KEY")
        except: 
            return None
    
    @staticmethod
    def extract_text_from_file(uploaded_file):
        text = ""
        try:
            if uploaded_file.name.endswith('.pdf'):
                pdf_reader = PyPDF2.PdfReader(uploaded_file)
                for page in pdf_reader.pages: 
                    extracted = page.extract_text()
                    if extracted: text += extracted + "\n"
            elif uploaded_file.name.endswith('.docx'):
                doc = Document(uploaded_file)
                for para in doc.paragraphs: text += para.text + "\n"
            else:
                text = uploaded_file.read().decode("utf-8")
        except Exception as e:
            st.error(f"Dosya okuma hatası: {e}")
        return text

    @staticmethod
    def generate_from_text(text, num_questions=3, provider="google", model_name='gemini-2.5-flash',target_course='AI-GEN'):
        api_key = AIGenerator.get_api_key(provider)
        
        if not api_key:
            st.warning(f"⚠️ {provider.title()} servisi için API anahtarı tanımlanmamış.", icon="🤖")
            return [] 
        
        response_schema = {
            "type": "ARRAY",
            "items": {
                "type": "OBJECT",
                "properties": {
                    "QuestionText": {"type": "STRING", "description": "Soru metni"},
                    "Options": {
                        "type": "OBJECT",
                        "description": "Çoktan seçmeli şıklar",
                        "properties": {
                            "A": {"type": "STRING"},
                            "B": {"type": "STRING"},
                            "C": {"type": "STRING"},
                            "D": {"type": "STRING"},
                            "E": {"type": "STRING"}
                        }
                    },
                    "CorrectAnswer": {"type": "STRING", "description": "Doğru cevap anahtarı (örn: A)"},
                    "Complexity": {"type": "INTEGER", "description": "Zorluk seviyesi (1, 2 veya 3)"},
                    "Score": {"type": "NUMBER", "description": "Soruya atanacak puan"}
                },
                "required": ["QuestionText", "Options", "CorrectAnswer", "Complexity", "Score"]
            }
        }
        
        generation_config = genai.GenerationConfig(
            response_mime_type="application/json",
            response_schema=response_schema
        )

        prompt = f"""
        Aşağıdaki metni analiz et ve {num_questions} adet akademik sınav sorusu oluştur.
        
        Metin: "{text[:8000]}"
        
        Kural: Soruların tipi Çoktan Seçmeli (MC) olsun. Çıktıyı tam olarak tanımlanan JSON şemasına göre oluştur.
        """
        
        try:
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel(model_name)
            response = model.generate_content(prompt, generation_config=generation_config)
            response_text = response.text
            response_text = response.text.replace("```json", "").replace("```", "").strip()
            questions_data = json.loads(response_text)
            
            final_questions = []
            for q in questions_data:
                final_questions.append({
                    'CourseCode': target_course, 
                    'TopicArea': 'AI Üretimi',
                    'Complexity': int(q.get('Complexity', 2)), 
                    'QuestionType': 'MC',
                    'Score': float(q.get('Score', 10)), 
                    'QuestionText': q.get('QuestionText', 'Soru metni alınamadı'),
                    'Options': q.get('Options', {}), 
                    'CorrectAnswer': q.get('CorrectAnswer', 'A'),
                    'CreatedBy': st.session_state['user']['Username']
                })
            return final_questions

        except json.JSONDecodeError as e:
            error_details = f"JSON format hatası. Model düzgün JSON döndüremedi. Hata: {str(e)}"
            if response_text:
                 error_details += f"\n\nModelin Ham Çıktısı (ilk 500 karakter): {response_text[:500]}..."
            st.error(f"AI İşlem Hatası: {error_details}")
            return []
        
        except api_exceptions.GoogleAPICallError as e: 
            error_message = e.message if hasattr(e, 'message') else str(e)
            
            error_msg = f"API Hatası (Gemini): {error_message}"
            
            if "API key not valid" in error_message:
                error_msg = "API Anahtarı Geçersiz veya Yetkisiz. Lütfen kontrol edin."
            elif "quota" in error_message.lower():
                error_msg = "Kota Aşıldı. Lütfen kullanım limitlerinizi kontrol edin."
            
            st.error(f"AI Genel Hata: {error_msg}")
            return []

        except Exception as e:
            st.error(f"AI Genel Hata: {str(e)}")
            st.caption("API anahtarınızın doğru olduğundan ve servis limitlerinizi aşmadığınızdan emin olun.")
            return []

    @staticmethod
    def analyze_question_bloom(question_text, provider="google", model_name='gemini-2.5-flash'):
        """Sorunun Bloom Taksonomisi seviyesini ve önerileri analiz eder."""
        api_key = AIGenerator.get_api_key(provider)
        if not api_key: return "API Key Eksik"
        
        prompt = f"""
        Aşağıdaki akademik sınav sorusunu eğitim bilimleri açısından analiz et.
        
        Soru: "{question_text}"
        
        Lütfen şu formatta JSON döndür:
        {{
            "bloom_level": "Bilgi/Kavrama/Uygulama/Analiz/Sentez/Değerlendirme",
            "reason": "Kısaca neden bu seviyede olduğu",
            "improvement_suggestion": "Soruyu daha üst seviyeye taşımak için kısa bir öneri"
        }}
        """
        try:
            genai.configure(api_key=api_key)
            model = genai.GenerativeModel(model_name)
            response = model.generate_content(prompt)
            return json.loads(response.text.replace("```json", "").replace("```", "").strip())
        except Exception as e:
            return {"error": str(e)}

class ExamPDFEngine(FPDF):
    def __init__(self, meta, is_answer_key=False, group_name="A", classical_lines=5):
        super().__init__()
        self.meta = meta
        self.is_answer_key = is_answer_key
        self.group_name = group_name
        self.classical_lines = classical_lines 
        check_and_download_font()
        self.set_auto_page_break(auto=True, margin=20)
        if os.path.exists(FONT_FILENAME):
            self.add_font('DejaVu', '', FONT_FILENAME, uni=True)
            self.add_font('DejaVu', 'B', FONT_FILENAME, uni=True)
            self.font_family = 'DejaVu'
        else:
            self.font_family = 'Arial' 

    def header(self):
        self.set_font(self.font_family, 'B', 14)
        title_suffix = " - CEVAP ANAHTARI" if self.is_answer_key else ""
        self.cell(0, 10, f"{self.meta['title']}{title_suffix}", 0, 1, 'C')
        self.set_font(self.font_family, '', 10)
        self.cell(0, 5, f"Ders: {self.meta['course']} | Grup: {self.group_name} | Tarih: {datetime.now().strftime('%d.%m.%Y')}", 0, 1, 'C')
        self.ln(5)
        
        if not self.is_answer_key:
            self.set_line_width(0.3)
            start_y = self.get_y()
            self.rect(10, start_y, 190, 25)
            self.set_font(self.font_family, '', 9)
            
            self.set_xy(12, start_y + 3)
            self.cell(20, 5, "Adı Soyadı:", 0, 0)
            self.cell(70, 5, "."*40, 0, 1)
            
            self.set_xy(12, start_y + 11)
            self.cell(20, 5, "Numarası:", 0, 0)
            self.cell(70, 5, "."*40, 0, 1)

            self.set_xy(110, start_y + 3)
            self.cell(15, 5, "Bölümü:", 0, 0)
            self.cell(70, 5, "."*40, 0, 1)
            
            self.set_xy(110, start_y + 11)
            self.cell(15, 5, "İmza:", 0, 0)
            self.cell(70, 5, "."*40, 0, 1)
            
            self.set_xy(12, start_y + 19)
            self.set_font(self.font_family, 'B', 12)
            self.cell(190, 5, f"KİTAPÇIK TÜRÜ: {self.group_name}", 0, 1, 'C')
            self.ln(5) 
        self.line(10, self.get_y(), 200, self.get_y())
        self.ln(8)

    def footer(self):
        self.set_y(-15)
        self.set_font(self.font_family, '', 8)
        self.cell(0, 10, f'SSOP Pro v5.2 - Grup {self.group_name} - Sayfa {self.page_no()}', 0, 0, 'C')

    def generate_content(self, questions):
        self.add_page()
        self.set_font(self.font_family, '', 11)
        
        for idx, q in enumerate(questions, 1):
            q_text = q['QuestionText']
            score_txt = f"({q['Score']} Puan)" if 'Score' in q else ""
            
            if self.is_answer_key:
                 header = f"{idx}. {q_text}\n   >>> DOĞRU CEVAP: {q.get('CorrectAnswer', '-')}"
                 self.set_text_color(200, 0, 0) 
            else:
                 header = f"{idx}. {q_text} {score_txt}"
                 self.set_text_color(0, 0, 0)

            if self.get_y() > 250: self.add_page()
            
            self.set_font(self.font_family, 'B', 11)
            self.multi_cell(0, 6, header)
            self.set_font(self.font_family, '', 11)
            self.set_text_color(0, 0, 0)
            self.ln(2)

            if not self.is_answer_key:
                if q['QuestionType'] == 'MC':
                    opts = q['Options'] if isinstance(q['Options'], dict) else (json.loads(q['Options']) if isinstance(q['Options'], str) else {})
                    if opts:
                        for k, v in sorted(opts.items()):
                            self.cell(10, 6, f"{k})", 0, 0)
                            self.multi_cell(0, 6, str(v))
                elif q['QuestionType'] == 'TF':
                    self.cell(5)
                    self.cell(30, 8, "◯ Doğru", 0, 0)
                    self.cell(30, 8, "◯ Yanlış", 0, 1)
                elif q['QuestionType'] == 'CL': 
                    self.ln(2)
                    for _ in range(self.classical_lines):
                        self.cell(0, 5, "_"*80, 0, 1) 
                    self.ln(3)

            self.ln(3)
            if not self.is_answer_key:
                self.set_draw_color(200,200,200)
                self.line(10, self.get_y(), 200, self.get_y())
                self.set_draw_color(0,0,0)
                self.ln(4)
            
    def get_pdf_bytes(self):
        return self.output(dest='S').encode('latin-1')

class ExamDocxEngine:
    """Word formatında sınav çıktısı üretir (Tablo Düzeni ile Geliştirildi)"""
    def __init__(self, meta, is_answer_key=False, group_name="A", classical_lines=5):
        self.meta = meta
        self.is_answer_key = is_answer_key
        self.group_name = group_name
        self.classical_lines = classical_lines
        self.doc = Document()
        
        section = self.doc.sections[0]
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)
        
    def generate(self, questions):
        h1 = self.doc.add_heading(self.meta['title'], 0)
        h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        info = f"Ders: {self.meta['course']} | Grup: {self.group_name} | Tarih: {datetime.now().strftime('%d.%m.%Y')}"
        p_info = self.doc.add_paragraph(info)
        p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        if not self.is_answer_key:
            self.doc.add_paragraph("\n")
            info_table = self.doc.add_table(rows=2, cols=2)
            info_table.autofit = True
            
            r1c1, r1c2 = info_table.rows[0].cells
            r2c1, r2c2 = info_table.rows[1].cells
            
            r1c1.text = "Adı Soyadı: ..........................................."
            r1c2.text = "Numarası: ............................"
            r2c1.text = "İmza: ................................................."
            r2c2.text = f"KİTAPÇIK TÜRÜ: {self.group_name}"
            
            r2c2.paragraphs[0].runs[0].bold = True
            
            self.doc.add_paragraph("\n" + "="*80 + "\n")
        
        self.doc.add_paragraph("")
        
        table = self.doc.add_table(rows=0, cols=2)
        table.autofit = False 
        table.columns[0].width = Inches(0.4) 
        table.columns[1].width = Inches(6.5)
        
        for idx, q in enumerate(questions, 1):
            row = table.add_row()
            cell_num = row.cells[0]
            cell_content = row.cells[1]
            
            cell_num.text = f"{idx}."
            cell_num.paragraphs[0].runs[0].bold = True
            
            q_text = q['QuestionText']
            score_txt = f"({q['Score']} Puan)"
            
            p = cell_content.paragraphs[0]
            runner = p.add_run(f"{q_text} ")
            runner.bold = True
            runner_score = p.add_run(score_txt)
            runner_score.font.size = Pt(9)
            
            if self.is_answer_key:
                ans_p = cell_content.add_paragraph()
                ans_runner = ans_p.add_run(f">> DOĞRU CEVAP: {q.get('CorrectAnswer', '-')}")
                ans_runner.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
                ans_runner.bold = True
            else:
                if q['QuestionType'] == 'MC':
                    opts = q['Options'] if isinstance(q['Options'], dict) else (json.loads(q['Options']) if isinstance(q['Options'], str) else {})
                    if opts:
                        for k, v in sorted(opts.items()):
                            p_opt = cell_content.add_paragraph(f"{k}) {v}")
                            p_opt.paragraph_format.space_after = Pt(2)
                            
                elif q['QuestionType'] == 'TF':
                    cell_content.add_paragraph("( ) Doğru    ( ) Yanlış")
                
                elif q['QuestionType'] == 'CL': 
                    cell_content.add_paragraph("") 
                    for _ in range(self.classical_lines):
                        p_line = cell_content.add_paragraph("_" * 90)
                        p_line.paragraph_format.space_after = Pt(0)

            cell_content.add_paragraph("")

    def get_docx_bytes(self):
        buffer = BytesIO()
        self.doc.save(buffer)
        return buffer.getvalue()
   
def show_question_edit_form(q_id, q_data_full):
    
    q_data = q_data_full.copy()
    type_map = {"MC": "Çoktan Seçmeli", "TF": "Doğru/Yanlış", "CL": "Klasik"}
    user = st.session_state['user']

    st.subheader(f"🛠️ Soru Düzenle (ID: {q_id} | Tip: {type_map.get(q_data['QuestionType'])})")
    
    with st.form(f"quick_edit_form_{q_id}"):
        ne_text = st.text_area("Soru Metni", value=q_data['QuestionText'], height=100)
        
        c_ne1, c_ne2, c_ne3, c_ne4 = st.columns(4)
        ne_course = c_ne1.text_input("Ders Kodu", value=q_data['CourseCode'])
        ne_topic = c_ne2.text_input("Konu Alanı", value=q_data['TopicArea'])
        ne_score = c_ne3.number_input("Puan", value=float(q_data['Score']), min_value=1.0)
        ne_diff = c_ne4.slider("Zorluk", 1, 3, int(q_data['Complexity']))
        
        ne_ans = st.text_input(f"Doğru Cevap ({'MC için (A,B,C..), TF için (Doğru/Yanlış), CL için boş bırakın'})", value=q_data.get('CorrectAnswer', ''))
        
        new_opts = None 
        if q_data['QuestionType'] == 'MC':
            current_opts = q_data.get('Options')
            if isinstance(current_opts, str):
                try: current_opts_dict = json.loads(current_opts)
                except: current_opts_dict = {}
            else:
                current_opts_dict = current_opts if current_opts else {}
                
            st.caption("Seçenekler (A: Şık1, B: Şık2 şeklinde, her satıra bir şık)")
            opts_text = "\n".join([f"{k}: {v}" for k, v in current_opts_dict.items()])
            ne_opts_text = st.text_area("Seçenekler", value=opts_text, height=100)
            
            try:
                new_opts = {}
                for line in ne_opts_text.split('\n'):
                    line = line.strip()
                    if not line: continue
                    if ':' in line:
                        key, value = line.split(':', 1)
                        if re.match(OPTION_KEY_REGEX, key.strip()):
                            new_opts[key.strip().upper()] = value.strip()
                        else:
                            st.error("Şık anahtarları tek bir harf (A, B, C...) olmalıdır.")
                            new_opts = None 
                            break
                    else:
                        st.error("Seçenekler 'Anahtar: Metin' formatında olmalıdır.")
                        new_opts = None
                        break
            except Exception as e:
                st.error(f"Seçenek format hatası: {e}")
                new_opts = None
        
        if st.form_submit_button("💾 Değişiklikleri Kaydet", type="primary"):
            if q_data['QuestionType'] == 'MC' and new_opts is None:
                st.error("Çoktan Seçmeli soru için şık formatını veya anahtarını düzeltin.")
                return

            update_payload = q_data.copy()
            update_payload['QuestionText'] = ne_text
            update_payload['CourseCode'] = ne_course
            update_payload['TopicArea'] = ne_topic
            update_payload['Score'] = ne_score
            update_payload['Complexity'] = ne_diff
            update_payload['CorrectAnswer'] = ne_ans
            update_payload['Options'] = new_opts if q_data['QuestionType'] == 'MC' else None
            
            if db.update_question(q_id, update_payload, user['Username']):
                 st.success("Soru başarıyla güncellendi!")
                 time.sleep(1)
                 st.session_state.pop('edit_qid', None) 
                 st.rerun()

def check_similarity(new_text, existing_questions, threshold=0.8):
    """Yeni metin ile mevcut sorular arasındaki benzerliği ölçer."""
    similar_questions = []
    if not existing_questions:
        return []
        
    for q in existing_questions:
        existing_text = q['QuestionText']
        ratio = difflib.SequenceMatcher(None, new_text.lower(), existing_text.lower()).ratio()
        
        if ratio > threshold:
            similar_questions.append({
                "ID": q['QuestionID'],
                "Text": existing_text,
                "Ratio": int(ratio*100)
            })
    return similar_questions

# ==============================================================================
# 4. YARDIMCI SINIFLAR (DBP ÇEKİCİ)
# ==============================================================================
@contextmanager
def get_db_connection():
    conn = sqlite3.connect(DB_FILE, check_same_thread=False, timeout=3.0) 
    conn.row_factory = sqlite3.Row
    try:
        yield conn
        conn.commit()
    except sqlite3.Error as e:
        conn.rollback()
        st.error(f"Veritabanı İşlem Hatası: {e}") 
        raise e
    except Exception as e:
        conn.rollback()
        st.error(f"Veritabanı Genel Hatası: {e}")
        raise e
    finally:
        conn.close()

class DatabaseManager:
    def __init__(self):
        self.init_db()

    def init_db(self):
        """
        Veritabanı tablolarını oluşturur. Fakülte ve Program sütunlarını kontrol eder/ekler.
        """
        with get_db_connection() as conn:
            cursor = conn.cursor()
            
            cursor.execute("PRAGMA table_info(users)")
            u_cols = [info[1] for info in cursor.fetchall()]

            if 'Theme' not in u_cols:
                 try: cursor.execute("ALTER TABLE users ADD COLUMN Theme TEXT DEFAULT 'Glassmorphism'")
                 except Exception: pass
            if 'Email' not in u_cols: 
                 try: cursor.execute("ALTER TABLE users ADD COLUMN Email TEXT")
                 except Exception: pass
            
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS users (
                    Username TEXT PRIMARY KEY,
                    Password TEXT NOT NULL,
                    Role TEXT DEFAULT 'Öğretim Üyesi',
                    FullName TEXT,
                    Photo TEXT,
                    Theme TEXT DEFAULT 'Glassmorphism',
                    Email TEXT 
                )
            """)
            
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS courses (
                    CourseCode TEXT PRIMARY KEY,
                    CourseName TEXT NOT NULL,
                    Faculty TEXT,    -- YENİ: Fakülte Bilgisi
                    Program TEXT,    -- YENİ: Bölüm/Program Bilgisi
                    CourseLevel TEXT,             
                    CourseType TEXT,              
                    CoursePeriod TEXT,            
                    LocalCredit REAL,             
                    AKTSCredit REAL,              
                    CourseLanguage TEXT,          
                    Coordinator TEXT,             
                    CourseGoal TEXT,              
                    CourseDesc TEXT,              
                    Prerequisites TEXT,           
                    CourseContent TEXT,           
                    LearningOutcomes TEXT,        
                    ProgramContribute TEXT,       
                    WorkloadDetails TEXT,         
                    EvaluationMethods TEXT,       
                    Resources TEXT,               
                    CreatedBy TEXT NOT NULL,
                    CreatedAt TIMESTAMP DEFAULT CURRENT_TIMESTAMP
                )
            """)

            cursor.execute("PRAGMA table_info(courses)")
            c_cols = [info[1] for info in cursor.fetchall()]
            
            if 'Program' not in c_cols:
                try: cursor.execute("ALTER TABLE courses ADD COLUMN Program TEXT")
                except Exception: pass
            
            if 'Faculty' not in c_cols:
                try: cursor.execute("ALTER TABLE courses ADD COLUMN Faculty TEXT")
                except Exception: pass

            NEW_ADMIN_EMAIL = 'ozdemirumut@gmail.com'
            if not cursor.execute("SELECT 1 FROM users WHERE Username = 'admin'").fetchone():
                cursor.execute("INSERT INTO users (Username, Password, Role, FullName, Theme, Email) VALUES (?, ?, ?, ?, ?, ?)", 
                               ('admin', hash_password('admin'), 'Admin', 'Sistem Yöneticisi', 'Glassmorphism', NEW_ADMIN_EMAIL))
            else:
                cursor.execute("UPDATE users SET Email = ? WHERE Username = 'admin'", (NEW_ADMIN_EMAIL,))
            
            cursor.execute("PRAGMA table_info(questions)")
            q_cols = [info[1] for info in cursor.fetchall()]

            if 'LastEditedBy' not in q_cols:
                try: cursor.execute("ALTER TABLE questions ADD COLUMN LastEditedBy TEXT")
                except Exception: pass
            if 'LastEditedAt' not in q_cols:
                try: cursor.execute("ALTER TABLE questions ADD COLUMN LastEditedAt TIMESTAMP")
                except Exception: pass
            if 'UsageCount' not in q_cols: 
                try: cursor.execute("ALTER TABLE questions ADD COLUMN UsageCount INTEGER DEFAULT 0")
                except Exception: pass
            
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS questions (
                    QuestionID INTEGER PRIMARY KEY AUTOINCREMENT,
                    CourseCode TEXT NOT NULL,
                    TopicArea TEXT NOT NULL,
                    Complexity INTEGER NOT NULL,
                    QuestionType TEXT NOT NULL,
                    Score REAL NOT NULL,
                    QuestionText TEXT NOT NULL,
                    Options TEXT,
                    CorrectAnswer TEXT,
                    CreatedBy TEXT NOT NULL,
                    CreatedAt TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    LastEditedBy TEXT,
                    LastEditedAt TIMESTAMP,
                    UsageCount INTEGER DEFAULT 0
                )
            """)
            
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS created_exams (
                    ExamID INTEGER PRIMARY KEY AUTOINCREMENT,
                    Title TEXT,
                    CourseCode TEXT,
                    TotalScore REAL,
                    ExamData TEXT,
                    CreatedBy TEXT,
                    CreatedAt TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    IsArchived INTEGER DEFAULT 0,
                    Status TEXT DEFAULT 'Final'
                )
            """)
            
            cursor.execute("""
                CREATE TABLE IF NOT EXISTS audit_logs (
                    LogID INTEGER PRIMARY KEY AUTOINCREMENT,
                    Timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                    Username TEXT,
                    Action TEXT, 
                    Details TEXT
                )
            """)
            
            cursor.execute("PRAGMA table_info(created_exams)")
            columns = [info[1] for info in cursor.fetchall()]
            if 'TotalScore' not in columns:
                try: cursor.execute("ALTER TABLE created_exams ADD COLUMN TotalScore REAL DEFAULT 0")
                except Exception: pass
            if 'IsArchived' not in columns:
                try: cursor.execute("ALTER TABLE created_exams ADD COLUMN IsArchived INTEGER DEFAULT 0")
                except Exception: pass
            if 'Status' not in columns:
                try: cursor.execute("ALTER TABLE created_exams ADD COLUMN Status TEXT DEFAULT 'Final'")
                except Exception: pass

    def log_action(self, username, action, details):
        """
        Kullanıcı aksiyonlarını veritabanına kaydeder.
        Bu fonksiyon, *get_db_connection* kullanılmadığı yerlerde (örn: init_db) veya 
        sadece loglama ihtiyacı olan yerlerde (kendi bağlantısını açıp kapatarak) kullanılmalıdır.
        """
        max_retries = 5 
        base_wait = 0.3 

        for attempt in range(max_retries):
            conn = None
            try:
                conn = sqlite3.connect(DB_FILE, check_same_thread=False, timeout=5.0) 
                
                conn.execute("INSERT INTO audit_logs (Username, Action, Details) VALUES (?, ?, ?)", 
                             (username, action, details))
                conn.commit()
                return 
            
            except sqlite3.OperationalError as e:
                if 'database is locked' in str(e).lower() and attempt < max_retries - 1:
                    wait_time = base_wait * (attempt + 1) + random.uniform(0, 0.1) 
                    time.sleep(wait_time) 
                    continue 
                else:
                    st.error(f"Kritik Loglama Hatası: {e}")
                    raise
            except Exception as e:
                 st.error(f"Kritik Loglama Genel Hatası: {e}")
                 raise
            finally:
                if conn:
                    conn.close()

    def create_user(self, username, password, fullname, role, theme, email):
        """
        Kullanıcı oluşturma ve loglama işlemini tek bir transaction'da yapar.
        """
        hashed_password = hash_password(password)
        with get_db_connection() as conn:
            try:
                conn.execute("INSERT INTO users (Username, Password, FullName, Role, Theme, Email) VALUES (?, ?, ?, ?, ?, ?)", 
                             (username, hashed_password, fullname, role, theme, email))
                
                log_username = st.session_state.get('user', {}).get('Username', 'System')
                log_details = f"New User: {username}, Role: {role}, Theme: {theme}, Email: {email}"
                
                conn.execute("INSERT INTO audit_logs (Username, Action, Details) VALUES (?, ?, ?)", 
                             (log_username, 'USER_CREATED', log_details))
                             
                return True
            except sqlite3.IntegrityError:
                return False 
            except Exception as e:
                st.error(f"Kullanıcı oluşturma hatası: {e}")
                return False

    def delete_user(self, username):
        """Kullanıcı silme fonksiyonu eklendi."""
        if username == 'admin': return False
        with get_db_connection() as conn:
            conn.execute("DELETE FROM users WHERE Username = ?", (username,))
            
            log_username = st.session_state.get('user', {}).get('Username', 'System')
            log_details = f"Deleted User: {username}"
            conn.execute("INSERT INTO audit_logs (Username, Action, Details) VALUES (?, ?, ?)", 
                         (log_username, 'USER_DELETED', log_details))
                         
        return True

    def login(self, username, password):
        with get_db_connection() as conn:
            user = conn.execute("SELECT * FROM users WHERE Username = ?", (username,)).fetchone()
            if user and check_password(password, user['Password']):
                return dict(user)
        return None
        
    def create_detailed_course(self, data, created_by):
        """
        Detaylı Ders Bilgi Paketi oluşturur. Fakülte, Program ve JSON dönüşümleri dahildir.
        """
        try:
            data['LocalCredit'] = float(data.get('LocalCredit', 0))
            data['AKTSCredit'] = float(data.get('AKTSCredit', 0))
        except ValueError:
             st.error("Kredi değerleri sayısal olmalıdır.")
             return False

        try:
             data['ProgramContribute'] = json.dumps(data.get('ProgramContribute', []), ensure_ascii=False)
             data['WorkloadDetails'] = json.dumps(data.get('WorkloadDetails', []), ensure_ascii=False)
             data['EvaluationMethods'] = json.dumps(data.get('EvaluationMethods', []), ensure_ascii=False)
             data['CourseContent'] = json.dumps(data.get('CourseContent', []), ensure_ascii=False)
             data['LearningOutcomes'] = json.dumps(data.get('LearningOutcomes', []), ensure_ascii=False)
        except Exception as e:
            st.error(f"Veri formatlama hatası (JSON): {e}")
            return False

        with get_db_connection() as conn:
            try:
                conn.execute("""
                    INSERT OR REPLACE INTO courses (
                        CourseCode, CourseName, Faculty, Program, CourseLevel, CourseType, CoursePeriod, 
                        LocalCredit, AKTSCredit, CourseLanguage, Coordinator, 
                        CourseGoal, CourseDesc, Prerequisites, CourseContent, 
                        LearningOutcomes, ProgramContribute, WorkloadDetails, 
                        EvaluationMethods, Resources, CreatedBy
                    ) VALUES (
                        :CourseCode, :CourseName, :Faculty, :Program, :CourseLevel, :CourseType, :CoursePeriod, 
                        :LocalCredit, :AKTSCredit, :CourseLanguage, :Coordinator, 
                        :CourseGoal, :CourseDesc, :Prerequisites, :CourseContent, 
                        :LearningOutcomes, :ProgramContribute, :WorkloadDetails, 
                        :EvaluationMethods, :Resources, :CreatedBy
                    )
                """, {**data, 'CreatedBy': created_by})
                
                fac = data.get('Faculty', '-')
                prog = data.get('Program', '-')
                conn.execute("INSERT INTO audit_logs (Username, Action, Details) VALUES (?, ?, ?)", 
                             (created_by, 'DBP_CREATED', f"Code: {data['CourseCode']}, Fac: {fac}, Prog: {prog}"))
                return True
            except sqlite3.IntegrityError:
                return False 
            except Exception as e:
                st.error(f"Veritabanı Kayıt Hatası: {e}")
                return False

    def get_courses(self, user_context):
        """Kullanıcının rolüne göre dersleri çeker."""
        query = "SELECT * FROM courses WHERE 1=1"
        params = []
        if user_context['Role'] != 'Admin':
            query += " AND CreatedBy = ?"
            params.append(user_context['Username'])
            
        query += " ORDER BY CourseCode ASC"
        with get_db_connection() as conn:
            return [dict(row) for row in conn.execute(query, params).fetchall()]

    def delete_course(self, course_code):
        """Dersi siler ve bu derse ait tüm soruları da siler."""
        with get_db_connection() as conn:
            conn.execute("DELETE FROM courses WHERE CourseCode = ?", (course_code,))
            conn.execute("DELETE FROM questions WHERE CourseCode = ?", (course_code,))

            log_username = st.session_state.get('user', {}).get('Username', 'System')
            log_details = f"Deleted Course: {course_code} and all associated questions."
            conn.execute("INSERT INTO audit_logs (Username, Action, Details) VALUES (?, ?, ?)", 
                         (log_username, 'COURSE_DELETED', log_details))
                         
        return True

    def _validate_mc_question(self, data):
        """MC sorunun şıklarını ve doğru cevabını kontrol eder."""
        if data['QuestionType'] == 'MC':
            options = data.get('Options')
            correct_answer = data.get('CorrectAnswer', '').strip().upper()

            if isinstance(options, str):
                try: options = json.loads(options)
                except: options = {}

            if not options or not correct_answer:
                return True 
            
            valid_keys = [k for k, v in options.items() if v and str(v).strip()]
            if correct_answer not in valid_keys:
                 raise ValueError(f"Doğru Cevap ('{correct_answer}') şıklar arasında bulunmuyor: {', '.join(valid_keys)}")
        return True

    def add_question(self, data):
        try:
             self._validate_mc_question(data)
        except ValueError as e:
             st.error(f"Soru Ekleme Hatası: {e}")
             return False

        options_json = json.dumps(data.get('Options', {}), ensure_ascii=False) if data.get('QuestionType') == 'MC' else None
        
        with get_db_connection() as conn:
            conn.execute("""
                INSERT INTO questions (CourseCode, TopicArea, Complexity, QuestionType, Score, QuestionText, Options, CorrectAnswer, CreatedBy)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
            """, (data['CourseCode'], data['TopicArea'], data['Complexity'], data['QuestionType'], data['Score'], 
                  data['QuestionText'], options_json, data['CorrectAnswer'], data['CreatedBy']))
            
            log_details = f"Course: {data['CourseCode']}, Topic: {data['TopicArea']}, Type: {data['QuestionType']}"
            conn.execute("INSERT INTO audit_logs (Username, Action, Details) VALUES (?, ?, ?)", 
                         (data['CreatedBy'], 'QUESTION_ADDED', log_details))
                         
        return True

    def update_question(self, q_id, data, editor_username):
        try:
             self._validate_mc_question(data)
        except ValueError as e:
             st.error(f"Soru Güncelleme Hatası: {e}")
             return False

        options_json = json.dumps(data.get('Options', {}), ensure_ascii=False) if data.get('QuestionType') == 'MC' else None
        
        with get_db_connection() as conn:
            conn.execute("""
                UPDATE questions SET CourseCode=?, TopicArea=?, Complexity=?, QuestionType=?, Score=?, QuestionText=?, Options=?, CorrectAnswer=?, LastEditedBy=?, LastEditedAt=?
                WHERE QuestionID=?
            """, (data['CourseCode'], data['TopicArea'], data['Complexity'], data['QuestionType'], data['Score'], 
                  data['QuestionText'], options_json, data['CorrectAnswer'], editor_username, datetime.now(), q_id))
            
            log_details = f"QID: {q_id}, Course: {data['CourseCode']}, Topic: {data['TopicArea']}"
            conn.execute("INSERT INTO audit_logs (Username, Action, Details) VALUES (?, ?, ?)", 
                         (editor_username, 'QUESTION_UPDATED', log_details))
                         
        return True

    def bulk_delete_questions(self, q_ids):
        if not q_ids: return
        with get_db_connection() as conn:
            ph = ','.join('?' for _ in q_ids)
            conn.execute(f"DELETE FROM questions WHERE QuestionID IN ({ph})", q_ids)
            
            log_username = st.session_state['user']['Username']
            log_details = f"Count: {len(q_ids)}, IDs: {q_ids[:5]}..."
            conn.execute("INSERT INTO audit_logs (Username, Action, Details) VALUES (?, ?, ?)", 
                         (log_username, 'QUESTIONS_DELETED', log_details))

    def get_questions(self, user_context, course_code=None):
        query = "SELECT * FROM questions WHERE 1=1"
        params = []
        if user_context['Role'] != 'Admin':
            query += " AND CreatedBy = ?"
            params.append(user_context['Username'])
        if course_code:
            query += " AND CourseCode = ?"
            params.append(course_code)
        query += " ORDER BY QuestionID DESC"
        with get_db_connection() as conn:
            return [dict(row) for row in conn.execute(query, params).fetchall()]
            
    def get_single_question(self, q_id):
         with get_db_connection() as conn:
             row = conn.execute("SELECT * FROM questions WHERE QuestionID = ?", (q_id,)).fetchone()
             return dict(row) if row else None

    def update_question_usage(self, q_ids):
        """Kullanılan soruların UsageCount sayacını artırır."""
        if not q_ids: return
        with get_db_connection() as conn:
             ph = ','.join('?' for _ in q_ids)
             conn.execute(f"UPDATE questions SET UsageCount = UsageCount + 1 WHERE QuestionID IN ({ph})", q_ids)

    def get_stats(self, user_context, course_code=None):
        where_clause = ""
        params = []
        
        if user_context['Role'] != 'Admin':
            where_clause = "WHERE CreatedBy = ?"
            params.append(user_context['Username'])
            
        if course_code:
            if where_clause:
                 where_clause += " AND CourseCode = ?"
            else:
                 where_clause = "WHERE CourseCode = ?"
            params.append(course_code)
            
        with get_db_connection() as conn:
            total = conn.execute(f"SELECT COUNT(*) FROM questions {where_clause}", params).fetchone()[0]
            courses = conn.execute(f"SELECT CourseCode, COUNT(*) as cnt FROM questions {where_clause} GROUP BY CourseCode", params).fetchall()
            topics = conn.execute(f"SELECT TopicArea, COUNT(*) as cnt FROM questions {where_clause} GROUP BY TopicArea", params).fetchall()
            avg_diff = conn.execute(f"SELECT AVG(Complexity) FROM questions {where_clause}", params).fetchone()[0] or 0
            
            exam_where = ""
            exam_params = []
            if user_context['Role'] != 'Admin':
                 exam_where = "WHERE CreatedBy = ?"
                 exam_params.append(user_context['Username'])
            exams = conn.execute(f"SELECT COUNT(*) FROM created_exams {exam_where}", exam_params).fetchone()[0]
            
            types = conn.execute(f"SELECT QuestionType, COUNT(*) FROM questions {where_clause} GROUP BY QuestionType", params).fetchall()
            diffs = conn.execute(f"SELECT Complexity, COUNT(*) FROM questions {where_clause} GROUP BY Complexity", params).fetchall()
            
            usage = conn.execute(f"SELECT QuestionID, UsageCount, QuestionText FROM questions {where_clause} ORDER BY UsageCount DESC LIMIT 5", params).fetchall()
            top_usage_list = [{"ID": row['QuestionID'], "Count": row['UsageCount'], "Text": row['QuestionText'][:60] + "..."} for row in usage]

            recent_questions = conn.execute(f"SELECT QuestionID, CreatedAt, QuestionText FROM questions {where_clause} ORDER BY CreatedAt DESC LIMIT 5", params).fetchall()
            recent_list = [{"ID": row['QuestionID'], "Date": row['CreatedAt'].split(' ')[0], "Text": row['QuestionText'][:60] + "..."} for row in recent_questions]

            return {
                'total': total, 'courses': dict(courses), 'topics': dict(topics), 'avg_diff': avg_diff, 
                'exams': exams, 'types': dict(types), 'diffs': dict(diffs), 'top_usage': top_usage_list,
                'recent_questions': recent_list
            }

    def save_exam(self, meta, questions, status='Final'):
        exam_json = json.dumps(questions, ensure_ascii=False)
        total_score = meta.get('score', sum(q['Score'] for q in questions if 'Score' in q))
        
        log_details = f"Title: {meta['title']}, Course: {meta['course']}, Score: {total_score}, QCount: {len(questions)}, Status: {status}"

        with get_db_connection() as conn:
            
            conn.execute("""
                INSERT INTO created_exams (Title, CourseCode, TotalScore, ExamData, CreatedBy, Status)
                VALUES (?, ?, ?, ?, ?, ?)
            """, (meta['title'], meta['course'], total_score, exam_json, meta['creator'], status))
            
            if status == 'Final':
                q_ids = [q['QuestionID'] for q in questions if 'QuestionID' in q]
                if q_ids:
                     ph = ','.join('?' for _ in q_ids)
                     conn.execute(f"UPDATE questions SET UsageCount = UsageCount + 1 WHERE QuestionID IN ({ph})", q_ids)
                 
            conn.execute("INSERT INTO audit_logs (Username, Action, Details) VALUES (?, ?, ?)", 
                         (meta['creator'], 'EXAM_CREATED', log_details))

            return conn.execute("SELECT last_insert_rowid()").fetchone()[0]

    def archive_exam(self, exam_id):
        """Sınavı veritabanında silmeden IsArchived=1 olarak işaretler."""
        with get_db_connection() as conn:
            conn.execute("UPDATE created_exams SET IsArchived = 1 WHERE ExamID = ?", (exam_id,))
            
            log_username = st.session_state['user']['Username']
            log_details = f"ExamID: {exam_id}"
            conn.execute("INSERT INTO audit_logs (Username, Action, Details) VALUES (?, ?, ?)", 
                         (log_username, 'EXAM_ARCHIVED', log_details))

    def get_exams(self, user_context, status=None, is_archived=0):
        """Sınavları çeker (is_archived parametresi eklendi)."""
        query = "SELECT * FROM created_exams WHERE IsArchived = ?"
        params = [is_archived]
        
        if user_context['Role'] != 'Admin':
            query += " AND CreatedBy = ?"
            params.append(user_context['Username'])
            
        if status:
            query += " AND Status = ?"
            params.append(status)
            
        query += " ORDER BY CreatedAt DESC"
        with get_db_connection() as conn:
            return [dict(row) for row in conn.execute(query, params).fetchall()]

    def get_single_exam(self, exam_id):
        with get_db_connection() as conn:
             row = conn.execute("SELECT * FROM created_exams WHERE ExamID = ?", (exam_id,)).fetchone()
             return dict(row) if row else None

    def get_all_users(self):
        with get_db_connection() as conn:
            return [dict(row) for row in conn.execute("SELECT Username, Role, FullName, Theme, Email FROM users").fetchall()]
    
    def get_audit_logs(self, limit=10):
        """Audit loglarını çeker."""
        with get_db_connection() as conn:
             return [dict(row) for row in conn.execute("SELECT * FROM audit_logs ORDER BY Timestamp DESC LIMIT ?", (limit,)).fetchall()]

    def reset_password(self, username, new_password):
        """Admin tarafından kullanıcı şifresini zorla sıfırlar."""
        if not new_password: return False
        
        hashed_pw = hash_password(new_password)
        with get_db_connection() as conn:
            conn.execute("UPDATE users SET Password = ? WHERE Username = ?", (hashed_pw, username))
            
            log_username = st.session_state.get('user', {}).get('Username', 'System')
            conn.execute("INSERT INTO audit_logs (Username, Action, Details) VALUES (?, ?, ?)", 
                         (log_username, 'PASSWORD_RESET', f"Reset password for: {username}"))
        return True
    
    def get_course_context_for_ai(self, course_code):
        """AI için dersin içeriğini ve çıktılarını metin olarak hazırlar."""
        with get_db_connection() as conn:
            row = conn.execute("SELECT CourseName, CourseContent, LearningOutcomes FROM courses WHERE CourseCode = ?", (course_code,)).fetchone()
            if not row:
                return None
            
            data = dict(row)
            context_text = f"Ders Kodu: {course_code}\nDers Adı: {data['CourseName']}\n\n"
            
            try:
                content_list = json.loads(data['CourseContent']) if data['CourseContent'] else []
                if content_list:
                    context_text += "HAFTALIK DERS İÇERİĞİ:\n"
                    for idx, item in enumerate(content_list, 1):
                        context_text += f"{idx}. Hafta: {item}\n"
            except:
                pass 
                
            context_text += "\n"
            
            try:
                outcomes_list = json.loads(data['LearningOutcomes']) if data['LearningOutcomes'] else []
                if outcomes_list:
                    context_text += "DERS ÖĞRENME ÇIKTILARI:\n"
                    for item in outcomes_list:
                        context_text += f"- {item}\n"
            except:
                pass
                
            return context_text

db = DatabaseManager()

class DBPFetcher:
    def __init__(self, db_manager):
        self.db = db_manager
        self.headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36'
        }
        self.base_url_check = "https://dbp.erciyes.edu.tr/Courses/Course.aspx"
        
    def fetch_course_data(self, url):
        """URL'den DBP verisini çeker ve ayrıştırır."""
        
        if not url.startswith(self.base_url_check):
            return {"error": "Hata: Geçersiz DBP URL'si. Erciyes Üniversitesi DBP formatı bekleniyor."}

        try:
            response = requests.get(url, headers=self.headers, timeout=10)
            response.raise_for_status() 
        except requests.exceptions.RequestException as e:
            return {"error": f"HTTP İstek Hatası: {e}"}

        soup = BeautifulSoup(response.content, 'html.parser')
        data = {}

        def get_text_by_label(label_text, default="Bilinmiyor"):
            label_tag = soup.find(lambda tag: tag.name in ['td', 'th'] and label_text in tag.get_text(strip=True))
            if label_tag:
                value_tag = label_tag.find_next_sibling('td')
                if value_tag:
                    return value_tag.get_text(strip=True)
            return default

        try:
            data['CourseName'] = soup.find('span', id=re.compile(r'lblCourseName')).get_text(strip=True) if soup.find('span', id=re.compile(r'lblCourseName')) else get_text_by_label("Dersin Adı", "Ders Adı Bulunamadı")
            data['CourseCode'] = soup.find('span', id=re.compile(r'lblCourseCode')).get_text(strip=True) if soup.find('span', id=re.compile(r'lblCourseCode')) else get_text_by_label("Dersin Kodu", "KOD_HATA")
            data['Coordinator'] = get_text_by_label("Koordinatör", "Koordinatör Yok")
            data['LocalCredit'] = float(get_text_by_label("Yerel Kredi", "0").replace(',', '.'))
            data['AKTSCredit'] = float(get_text_by_label("AKTS Kredisi", "0").replace(',', '.'))
            data['CourseType'] = get_text_by_label("Ders Tipi")
            data['CourseLevel'] = get_text_by_label("Ders Seviyesi")
            data['CourseLanguage'] = get_text_by_label("Eğitim Dili")
            data['CoursePeriod'] = get_text_by_label("Dönemi")
            data['CourseGoal'] = get_text_by_label("Dersin Amacı", "Amaç Metni Yok")
            data['CourseDesc'] = get_text_by_label("Dersin Tanımı", "Tanım Metni Yok")
            data['Prerequisites'] = get_text_by_label("Ön Koşul Dersleri", "Yok")
            data['LearningOutcomes'] = self._extract_list_by_header(soup, "Ders Öğrenme Çıktıları", re.compile(r'ulLearningOutcomes|repLearningOutcomes|lblLearningOutcomes'))
            data['Resources'] = self._extract_text_by_header(soup, "Önerilen/Zorunlu Kaynaklar", re.compile(r'lblResources|lblRecommendedResources'))
            data['EvaluationMethods'] = self._extract_evaluation_methods(soup, "Değerlendirme Sistemi")
            data['WorkloadDetails'] = self._extract_workload_details(soup, "AKTS/İş Yükü")
            data['ProgramContribute'] = self._extract_program_contribute(soup, "Program Yeterliliklerine Katkı")
            
            return data

        except Exception as e:
            st.exception(f"Veri ayrıştırma sırasında genel hata oluştu: {e}")
            return {"error": f"Veri ayrıştırma sırasında genel hata oluştu. Sayfa yapısı değişmiş olabilir: {e}"}

    def _extract_list_by_header(self, soup, header_text, regex_id):
        """Belirli bir başlığın altındaki liste öğelerini (li/td) veya özel ID'li alandan metni çeker."""
        target_tag = soup.find('span', id=regex_id)
        if target_tag:
            list_items = target_tag.find_all(['li', 'div']) 
            if list_items:
                return [item.get_text(strip=True) for item in list_items if item.get_text(strip=True)]
            return [line.strip() for line in target_tag.get_text(separator='\n').split('\n') if line.strip()]
        
        header = soup.find('h3', string=lambda s: s and header_text in s)
        if header:
            next_table = header.find_next('table')
            if next_table:
                rows = next_table.find_all('tr')
                items = []
                for row in rows[1:]: 
                    cols = row.find_all(['td', 'th'])
                    if len(cols) >= 2:
                        items.append(cols[1].get_text(strip=True))
                return items
        
        return ["Otomatik Çekim Başarısız: " + header_text]

    def _extract_text_by_header(self, soup, header_text, regex_id):
        """Belirli bir ID'ye sahip alandan veya etiketten metni çeker."""
        target_tag = soup.find('span', id=regex_id)
        if target_tag:
            return target_tag.get_text(strip=True, separator='\n')
        
        return "Otomatik Çekim Başarısız: " + header_text
        
    def _extract_evaluation_methods(self, soup, header_text):
        """Değerlendirme yöntemleri tablosundan veriyi çeker (örnek yapı)."""
        methods = []
        try:
            table = soup.find('span', id=re.compile(r'lblEvaluationMethods|repEvaluationMethods')).find_parent('table')
            if table:
                rows = table.find_all('tr')
                for row in rows[1:]: 
                    cols = row.find_all(['td', 'th'])
                    if len(cols) >= 3:
                        methods.append({
                            "method": cols[0].get_text(strip=True),
                            "count": int(cols[1].get_text(strip=True).replace('.', '').split(' ')[0]), 
                            "percent": int(cols[2].get_text(strip=True).replace('%', '').replace('.', '').split(' ')[0]),
                        })
        except:
             methods.append({"method": "Otomatik Çekim Hatası", "count": 0, "percent": 0})
             
        return methods
    
    def _extract_workload_details(self, soup, header_text):
        """İş yükü detayları tablosundan veriyi çeker (örnek yapı)."""
        workloads = []
        try:
            table = soup.find('span', id=re.compile(r'lblWorkloadDetails|repWorkloadDetails')).find_parent('table')
            if table:
                rows = table.find_all('tr')
                for row in rows[1:]:
                    cols = row.find_all(['td', 'th'])
                    if len(cols) >= 4:
                         workloads.append({
                            "activity": cols[0].get_text(strip=True),
                            "count": int(cols[1].get_text(strip=True).replace('.', '').split(' ')[0]),
                            "duration": int(cols[2].get_text(strip=True).replace('.', '').split(' ')[0]),
                            "total": int(cols[3].get_text(strip=True).replace('.', '').split(' ')[0]),
                        })
        except:
             workloads.append({"activity": "Otomatik Çekim Hatası", "count": 0, "duration": 0, "total": 0})
             
        return workloads

    def _extract_program_contribute(self, soup, header_text):
        """Program yeterliliklerine katkı tablosundan veriyi çeker (örnek yapı)."""
        contributes = []
        try:
            table = soup.find('span', id=re.compile(r'lblProgramContribute|repProgramContribute')).find_parent('table')
            if table:
                rows = table.find_all('tr')
                for row in rows[1:]:
                    cols = row.find_all(['td', 'th'])
                    if len(cols) >= 2:
                        level_text = cols[-1].get_text(strip=True)
                        level = 0
                        if '5' in level_text: level = 5
                        elif '4' in level_text: level = 4
                        elif '3' in level_text: level = 3
                        elif '2' in level_text: level = 2
                        elif '1' in level_text: level = 1
                        
                        contributes.append({
                            "text": cols[0].get_text(strip=True),
                            "level": level
                        })
        except:
             contributes.append({"text": "Otomatik Çekim Hatası", "level": 0})
             
        return contributes
        
dbp_fetcher = DBPFetcher(db)

# ==============================================================================
# 5. SAYFALAR
# ==============================================================================
def login_page():
    c1, c2, c3 = st.columns([1,1,1])
    with c2:
        st.markdown("<br><br>", unsafe_allow_html=True)
        st.markdown("""
        <div style='text-align: center; background: white; padding: 40px; border-radius: 15px; box-shadow: 0 10px 25px rgba(0,0,0,0.1);'>
            <h1 style='color: #2c3e50; margin-bottom:0;'>🎓 SSOP Pro</h1>
            <p style='color: #3498db; font-size: 1.1em; font-weight:bold;'>Enterprise Edition v5.2</p>
            <p style='color: #7f8c8d; font-size: 0.9em;'>Akademik Sınav Sistemi</p>
            <hr style='margin: 20px 0;'>
        </div>
        """, unsafe_allow_html=True)
        
        with st.form("login_form"):
            user = st.text_input("Kullanıcı Adı")
            pwd = st.text_input("Şifre", type="password")
            if st.form_submit_button("Giriş Yap", type="primary", use_container_width=True):
                account = db.login(user, pwd)
                if account:
                    st.session_state['user'] = account
                    st.session_state['user_current_theme'] = account.get('Theme', 'Glassmorphism')
                    st.rerun()
                else:
                    st.error("Hatalı kullanıcı adı veya şifre!")

def dashboard_page():
    user = st.session_state['user']
    st.markdown(f"## 👋 Hoşgeldin, {user['FullName']}")
    
    # Ders Seçimi
    all_courses = db.get_courses(user)
    course_options = ["Tüm Dersler"] + [c['CourseCode'] for c in all_courses]
    
    selected_course_filter = st.selectbox("📊 Gösterge Paneli Ders Filtresi", course_options, key="dashboard_course_filter")

    filter_code = selected_course_filter if selected_course_filter != "Tüm Dersler" else None
    
    stats = db.get_stats(user, course_code=filter_code)

    if selected_course_filter != "Tüm Dersler":
        st.subheader(f"Ders: {selected_course_filter}")
        
    if stats['total'] == 0:
        st.warning(f"Seçilen filtreye ({selected_course_filter}) göre görüntülenecek veri yok.")
        if selected_course_filter == "Tüm Dersler" and not all_courses:
             st.info("Sistemde hiç ders kaydı bulunmamaktadır. Lütfen 'Dersler' menüsünden ders ekleyin.")
        return
    
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Toplam Soru", stats['total'], "Adet")
    c2.metric("Aktif Ders", len(stats['courses']), "Ders")
    c3.metric("Ort. Zorluk", f"{stats['avg_diff']:.1f}", "/ 3.0")
    c4.metric("Üretilen Sınav", stats['exams'], "Adet")
    
    st.divider()
    
    col_g1, col_g2 = st.columns([2, 1])
    with col_g1:
        st.subheader("📊 Ders & Konu Dağılımı")
        
        if stats['courses']:
            treemap_data = []
            for course, count in stats['courses'].items():
                treemap_data.append(dict(Ders=course, Parent="Tüm Dersler", Soru=count))
            
            fig = px.treemap(
                treemap_data,
                path=['Parent', 'Ders'],
                values='Soru',
                color='Soru',
                color_continuous_scale='Blues'
            )
            st.plotly_chart(fig, use_container_width=True)
        else: st.info("Görüntülenecek ders verisi yok.")
            
    with col_g2:
        st.subheader("🧩 Soru Analizi")
        if stats.get('types') or stats.get('diffs'):
            df_t = pd.DataFrame(list(stats['types'].items()), columns=['Tip', 'Adet'])
            fig3 = px.pie(df_t, values='Adet', names='Tip', hole=0.6, title="Soru Tip Dağılımı")
            st.plotly_chart(fig3, use_container_width=True)
            
            df_d = pd.DataFrame(list(stats['diffs'].items()), columns=['Zorluk', 'Adet']).sort_values('Zorluk')
            fig4 = px.bar(df_d, x='Zorluk', y='Adet', title="Zorluk Seviyesi Dağılımı", color='Zorluk', text='Adet')
            st.plotly_chart(fig4, use_container_width=True)
        else: st.info("Analiz edilecek soru verisi yok.")
        
    st.divider()

    st.divider()
    st.subheader("📈 Zaman İçinde Soru Üretim Analizi (Time Series)")
    
    all_q_stats = db.get_questions(user, course_code=filter_code)
    if all_q_stats:
        df_ts = pd.DataFrame(all_q_stats)
        
        df_ts['CreatedAt'] = pd.to_datetime(df_ts['CreatedAt'])
        df_ts['Date'] = df_ts['CreatedAt'].dt.date
        
        daily_counts = df_ts.groupby('Date').size().reset_index(name='Soru Sayısı')
        
        daily_counts['Toplam Birikim'] = daily_counts['Soru Sayısı'].cumsum()
        
        tab_daily, tab_cum = st.tabs(["Günlük Aktivite", "Kümülatif Büyüme"])
        
        with tab_daily:
            fig_ts = px.bar(daily_counts, x='Date', y='Soru Sayısı', 
                            title="Günlük Soru Giriş Grafiği",
                            color='Soru Sayısı', color_continuous_scale='Viridis')
            st.plotly_chart(fig_ts, use_container_width=True)
            
        with tab_cum:
            fig_cum = px.line(daily_counts, x='Date', y='Toplam Birikim', 
                              title="Soru Bankası Büyüme Hızı",
                              markers=True, line_shape='spline') # Spline ile daha yumuşak hatlar
            fig_cum.add_scatter(x=daily_counts['Date'], 
                                y=daily_counts['Toplam Birikim'].rolling(window=7).mean(), 
                                mode='lines', name='7 Günlük Hareketli Ort.', 
                                line=dict(dash='dot', color='red'))
            st.plotly_chart(fig_cum, use_container_width=True)
    else:
        st.info("Zaman analizi için yeterli veri yok.")

    col_dash_1, col_dash_2 = st.columns(2)
    
    with col_dash_1:
        st.subheader("🔥 En Çok Kullanılan Sorular (Top 5)")
        if stats['top_usage']:
            df_top = pd.DataFrame(stats['top_usage'])
            df_top.rename(columns={'ID': 'Soru ID', 'Count': 'Kullanım Sayısı', 'Text': 'Soru Metni'}, inplace=True)
            st.table(df_top)
        else:
            st.info("Kullanım verisi bulunmamaktadır.")

    with col_dash_2:
        st.subheader("🆕 Son Eklenen Sorular (Top 5)")
        if stats['recent_questions']:
            df_recent = pd.DataFrame(stats['recent_questions'])
            df_recent.rename(columns={'ID': 'Soru ID', 'Date': 'Tarih', 'Text': 'Soru Metni'}, inplace=True)
            st.table(df_recent)
        else:
            st.info("Yakın zamanda eklenen soru bulunmamaktadır.")
            
    if user['Role'] == 'Admin':
        st.divider()
        st.subheader("🚨 Admin Aksiyon Logları (Son 10)")
        logs = db.get_audit_logs(limit=10)
        df_logs = pd.DataFrame(logs)
        if not df_logs.empty:
            df_logs['Timestamp'] = df_logs['Timestamp'].apply(lambda x: datetime.strptime(x.split('.')[0], '%Y-%m-%d %H:%M:%S').strftime('%Y-%m-%d %H:%M:%S'))
            st.dataframe(df_logs[['Timestamp', 'Username', 'Action', 'Details']], use_container_width=True, hide_index=True)
        else:
            st.info("Log kaydı bulunamadı.")

def course_management_page():
    user = st.session_state['user']
    st.title("📚 Ders Yönetimi Merkezi (DBP)")

    tab_create, tab_list, tab_fetch = st.tabs(["✍️ Ders Ekle", "📖 Kayıtlı Dersler", "🔗 DBP/Linkten Çek"])
    
    if 'workload_inputs' not in st.session_state:
        st.session_state['workload_inputs'] = {}
        defaults = {
            "Yüz yüze eğitim": (15, 2), 
            "Sınıf dışı ders çalışma süresi (ön çalışma, pekiştirme)": (10, 1),
            "Ara sınavlara hazırlık": (1, 8),
            "Ara sınavlar": (1, 1),
            "Yarıyıl sonu sınavına hazırlık": (1, 15),
            "Yarıyıl sonu sınavı": (1, 1),
        }
        for activity in ["Yüz yüze eğitim", "Sınıf dışı ders çalışma süresi (ön çalışma, pekiştirme)", "Ödevler", "Sunum / Seminer hazırlama", "Kısa sınavlar", "Ara sınavlara hazırlık", "Ara sınavlar", "Proje (Yarıyıl ödevi)", "Laboratuvar", "Arazi çalışması", "Yarıyıl sonu sınavına hazırlık", "Yarıyıl sonu sınavı", "Araştırma"]:
            count, duration = defaults.get(activity, (0, 0))
            st.session_state['workload_inputs'][f"wl_count_{activity}"] = count
            st.session_state['workload_inputs'][f"wl_duration_{activity}"] = duration

    if 'eval_inputs' not in st.session_state:
        st.session_state['eval_inputs'] = {}
        st.session_state['eval_inputs']["eval_count_Ara sınav"] = 1
        st.session_state['eval_inputs']["eval_percent_Ara sınav"] = 40
        st.session_state['eval_inputs']["eval_count_Yarıyıl sonu sınavı"] = 1
        st.session_state['eval_inputs']["eval_percent_Yarıyıl sonu sınavı"] = 60
    
    # ----------------------------------------------------
    # 1. SEKME: Detaylı Ders Ekleme Formu
    # ----------------------------------------------------
    with tab_create:
        st.subheader("➕ Yeni Ders Bilgi Paketi Oluştur")
        st.info("Bu form, Ders Bilgi Paketi'ndeki tüm detayları sekmeli yapıda kaydeder.")
        
        with st.form("add_detailed_course_form", clear_on_submit=False): 
            
            tab_temel, tab_icerik, tab_yeterlilik, tab_akts_deger = st.tabs([
                "1. Temel Bilgiler", 
                "2. İçerik ve Çıktılar", 
                "3. Yeterlilik Katkısı", 
                "4. İş Yükü & Değerlendirme " 
            ])
            
            with tab_temel:
                st.markdown("#### 1. Temel Ders Bilgileri")
                
                col_fac, col_prog = st.columns(2)
                new_faculty = col_fac.text_input("Fakülte / Yüksekokul (Örn: Mühendislik Fak.)", key="nc_faculty")
                new_program = col_prog.text_input("Bölüm / Program (Örn: Bilgisayar Müh.)", key="nc_program")

                c1, c2, c3 = st.columns(3)
                new_code = c1.text_input("Ders Kodu (Örn: SEC114)", key="nc_code").upper()
                new_name = c2.text_input("Ders Adı (Örn: MİKROBİYOLOJİ)", key="nc_name")
                new_level = c3.text_input("Ders Seviyesi", value="Birinci Düzey", key="nc_level")
                
                c4, c5, c6, c7 = st.columns(4)
                new_type = c4.selectbox("Ders Tipi", ["Seçmeli", "Zorunlu"], key="nc_type")
                new_period = c5.text_input("Ders Dönemi", value="1", key="nc_period")
                new_local_cr = c6.number_input("Yerel Kredi", min_value=0.0, value=3.0, key="nc_lcredit", format="%.2f")
                new_akts_cr = c7.number_input("AKTS Kredisi (Manuel Giriş)", min_value=0.0, value=3.0, key="nc_akts", format="%.2f") 
                
                new_coord = st.text_input("Koordinatör / Dersi Veren Öğretim Elemanı", value=user['FullName'], key="nc_coord")
                new_lang = st.selectbox("Eğitim Dili", ["Türkçe", "İngilizce"], key="nc_lang")
                
                st.markdown("---")
                st.markdown("#### 2. Ders Tanımı, Amaç ve Ön Koşullar")
                new_prereq = st.text_input("Ön Koşul Dersleri", value="Ön koşul yok", key="nc_prereq")
                new_goal = st.text_area("Dersin Amacı", placeholder="Dersin amacı...", key="nc_goal")
                new_desc = st.text_area("Dersin Tanımı", placeholder="Dersin içeriği hakkında kısa bilgi...", key="nc_desc")
                new_res = st.text_area("8. Önerilen veya Zorunlu Kaynaklar (Metin)", placeholder="Ders Kitabı: ...", key="nc_res")
                new_res_files = st.file_uploader("📂 Kaynak Dosyası Yükle (PDF, PPTX, DOCX)", accept_multiple_files=True, key="nc_res_files")

            with tab_icerik:
                st.markdown("#### 3. Ders İçeriği")
                course_content_list = []
                for i in range(1, 21):
                    content = st.text_input(f"{i}. Hafta/Konu", key=f"cc_week_{i}", value=f"Konu Metni {i}" if i < 3 else "")
                    if content.strip(): 
                        course_content_list.append(content.strip())
                
                st.markdown("---")
                st.markdown("#### 4. Ders Öğrenme Çıktıları")
                lo_list = []
                for i in range(1, 11):
                    lo = st.text_area(f"Çıktı {i}", key=f"lo_{i}", height=50, value=f"{i}. öğrenme çıktısı" if i <= 2 else "")
                    if lo.strip(): lo_list.append(lo.strip())

            with tab_yeterlilik:
                st.markdown("#### 5. Dersin Program Yeterliliklerine Katkı Seviyesi")
                pc_list = []
                col_pc_h1, col_pc_h2 = st.columns([4, 2])
                col_pc_h1.markdown("**Program Yeterliliği**")
                col_pc_h2.markdown("**Katkı Seviyesi (1-5 Yıldız)**")

                for i in range(1, 21):
                    col_pc_text, col_pc_star = st.columns([4, 2])
                    default_pc_text = ""
                    default_star = 0
                    pc_text = col_pc_text.text_input(f"Yeterlilik {i}", key=f"pc_text_{i}", value=default_pc_text, label_visibility="collapsed")
                    pc_star = col_pc_star.slider(f"Katkı Seviyesi {i}", min_value=0, max_value=5, value=default_star, key=f"pc_star_{i}", label_visibility="collapsed")
                    if pc_text.strip():
                        pc_list.append({"text": pc_text.strip(), "level": pc_star})
            
            with tab_akts_deger:
                st.markdown("#### 6. Öğrenim Faaliyetleri ve AKTS İşyükü")
                workload_activities = [
                    "Yüz yüze eğitim", "Sınıf dışı ders çalışma süresi (ön çalışma, pekiştirme)",
                    "Ödevler", "Sunum / Seminer hazırlama", "Kısa sınavlar", "Ara sınavlara hazırlık",
                    "Ara sınavlar", "Proje (Yarıyıl ödevi)", "Laboratuvar", "Arazi çalışması",
                    "Yarıyıl sonu sınavına hazırlık", "Yarıyıl sonu sınavı", "Araştırma"
                ]
                workload_data = [] 
                total_workload = 0
                
                w_c1, w_c2, w_c3, w_c4 = st.columns([3, 1, 1, 1])
                w_c1.write("**Faaliyet**")
                w_c2.write("**Sayısı**")
                w_c3.write("**Süresi**")
                w_c4.write("**Toplam**")

                for idx, activity in enumerate(workload_activities):
                    col_name, col_count, col_duration, col_total = st.columns([3, 1, 1, 1])
                    col_name.write(activity)
                    count_key = f"wl_count_{activity}"
                    duration_key = f"wl_duration_{activity}"
                    count = col_count.number_input(count_key, min_value=0, step=1, key=count_key, label_visibility="collapsed")
                    duration = col_duration.number_input(duration_key, min_value=0, step=1, key=duration_key, label_visibility="collapsed")
                    total = count * duration
                    col_total.write(f"**{total}**")
                    total_workload += total
                    workload_data.append({"activity": activity, "count": count, "duration": duration, "total": total})
                
                st.markdown("---")
                akts_calc_int = (total_workload + 29) // 30 if total_workload > 0 else 0
                st.markdown(f"**Toplam İş Yükü:** {total_workload} | **Hesaplanan AKTS:** {akts_calc_int}")
                
                st.markdown("---")
                st.markdown("#### 7. Değerlendirme Yöntemleri")
                eval_methods = ["Ara sınav", "Kısa sınav", "Ödev", "Yarıyıl sonu sınavı"]
                eval_data = [] 
                total_eval_percent = 0
                
                e_c1, e_c2, e_c3 = st.columns([3, 1, 1])
                e_c1.write("**Değerlendirme**")
                e_c2.write("**Sayısı**")
                e_c3.write("**Katkı Yüzdesi (%)**")
                
                for idx, method in enumerate(eval_methods):
                    col_name, col_count, col_percent = st.columns([3, 1, 1])
                    col_name.write(method)
                    count_key = f"eval_count_{method}"
                    percent_key = f"eval_percent_{method}"
                    count = col_count.number_input(count_key, min_value=0, max_value=20, step=1, key=count_key, label_visibility="collapsed")
                    percent = col_percent.number_input(percent_key, min_value=0, max_value=100, step=5, key=percent_key, label_visibility="collapsed")
                    total_eval_percent += percent
                    eval_data.append({"method": method, "count": count, "percent": percent})

                if total_eval_percent != 100:
                    st.warning("⚠️ Değerlendirme Katkı Yüzdesi toplamı 100 olmalıdır!")
            
            submit_button = st.form_submit_button("Detaylı DBP Kaydet", type="primary")

            if submit_button:
                if new_code and new_name:
                    final_resources_data = new_res 
                    if new_res_files or new_res:
                        saved_file_paths = []
                        if new_res_files:
                            upload_dir = "course_materials"
                            if not os.path.exists(upload_dir):
                                os.makedirs(upload_dir)
                            for uploaded_file in new_res_files:
                                safe_filename = f"{new_code}_{os.path.basename(uploaded_file.name)}".replace(" ", "_")
                                file_path = os.path.join(upload_dir, safe_filename)
                                with open(file_path, "wb") as f:
                                    f.write(uploaded_file.getbuffer())
                                saved_file_paths.append(file_path)
                        resource_json = {"text": new_res, "files": saved_file_paths}
                        final_resources_data = json.dumps(resource_json, ensure_ascii=False)

                    course_data = {
                        'CourseCode': new_code, 
                        'CourseName': new_name, 
                        'Faculty': new_faculty, 
                        'Program': new_program,  
                        'CourseLevel': new_level,
                        'CourseType': new_type, 'CoursePeriod': new_period, 'LocalCredit': new_local_cr,
                        'AKTSCredit': new_akts_cr, 'CourseLanguage': new_lang, 'Coordinator': new_coord,
                        'CourseGoal': new_goal, 'CourseDesc': new_desc, 'Prerequisites': new_prereq,
                        'Resources': final_resources_data, 
                        'CourseContent': course_content_list, 
                        'LearningOutcomes': lo_list,
                        'ProgramContribute': pc_list, 
                        'WorkloadDetails': workload_data, 
                        'EvaluationMethods': eval_data, 
                    }

                    if db.create_detailed_course(course_data, user['Username']):
                        st.toast(f"Ders Bilgi Paketi '{new_code}' başarıyla eklendi!", icon="✅")
                        st.session_state['course_rerun'] = True
                    else:
                        st.error(f"Ders Kodu ('{new_code}') zaten var veya bir veritabanı hatası oluştu.")
                else:
                    st.error("Ders Kodu ve Ders Adı zorunludur.")

    # ----------------------------------------------------
    # 2. SEKME: Kayıtlı Dersler Listesi
    # ----------------------------------------------------
    with tab_list:
        st.subheader("📖 Kayıtlı Dersler ve Detaylı Görünüm")
        if st.session_state.get('course_rerun'):
            st.session_state.pop('course_rerun')
            st.rerun() 
            
        courses = db.get_courses(user)
        if not courses:
            st.info("Henüz eklenmiş bir ders bulunmamaktadır.")
        else:
            for row in courses:
                with st.container(border=True):
                    c_header, c_delete = st.columns([6, 1])
                    
                    course_title = f"{row['CourseCode']} - {row['CourseName']}"
                    prog_info = row.get('Program') if row.get('Program') else "Belirtilmemiş"
                    fac_info = row.get('Faculty') if row.get('Faculty') else "Belirtilmemiş"

                    c_header.markdown(f"### 📘 {course_title}")
                    c_header.caption(f"**Fakülte:** {fac_info} | **Bölüm:** {prog_info} | **AKTS:** {row.get('AKTSCredit')}")
                    
                    if c_delete.button("🗑️ Sil", key=f"del_course_{row['CourseCode']}", type="primary"):
                        if db.delete_course(row['CourseCode']):
                            st.toast(f"Ders '{row['CourseCode']}' silindi!", icon="✅")
                            st.session_state['course_rerun'] = True
                            st.rerun()

                    with st.expander("📄 DERS BİLGİ PAKETİ DETAYLARINI GÖR", expanded=False):
                        display_data = dict(row)
                        for key in ['CourseContent', 'LearningOutcomes', 'ProgramContribute', 'WorkloadDetails', 'EvaluationMethods']:
                            if display_data.get(key) and isinstance(display_data[key], str):
                                try: display_data[key] = json.loads(display_data[key])
                                except: pass

                        d_tab1, d_tab2, d_tab3, d_tab4 = st.tabs(["Genel Bilgiler", "İçerik & Çıktılar", "İş Yükü & Değerlendirme", "Program Yeterliliği"])
                        
                        with d_tab1:
                            col_g1, col_g2 = st.columns(2)
                            col_g1.info(f"**Koordinatör:** {display_data.get('Coordinator', '-')}")
                            col_g1.write(f"**Eğitim Dili:** {display_data.get('CourseLanguage', '-')}")
                            
                            col_g2.success(f"**Kredi:** {display_data.get('LocalCredit', '-')} | **AKTS:** {display_data.get('AKTSCredit', '-')}")
                            col_g2.write(f"**Ön Koşul:** {display_data.get('Prerequisites', '-')}")
                            
                            st.markdown("---")
                            st.markdown(f"**Fakülte:** {fac_info}")
                            st.markdown(f"**Bölüm/Program:** {prog_info}")
                            st.markdown("---")
                            st.markdown("#### 🎯 Dersin Amacı")
                            st.write(display_data.get('CourseGoal', 'Belirtilmemiş'))
                            st.markdown("#### 📝 Dersin Tanımı")
                            st.write(display_data.get('CourseDesc', 'Belirtilmemiş'))
                            
                            st.markdown("#### 📚 Kaynaklar")
                            raw_resources = display_data.get('Resources', '')
                            try:
                                res_obj = json.loads(raw_resources)
                                if isinstance(res_obj, dict):
                                    if res_obj.get('text'): st.write(res_obj.get('text'))
                                    files = res_obj.get('files', [])
                                    if files:
                                        st.markdown("**📂 İndirilebilir Dosyalar:**")
                                        for f_path in files:
                                            if os.path.exists(f_path):
                                                with open(f_path, "rb") as file:
                                                    st.download_button(label=f"⬇️ İndir: {os.path.basename(f_path)}", data=file, file_name=os.path.basename(f_path), key=f"dl_btn_{f_path}")
                                else: st.write(raw_resources)
                            except: st.write(raw_resources if raw_resources else 'Belirtilmemiş')

                        with d_tab2:
                            col_content, col_outcomes = st.columns(2)
                            with col_content:
                                st.markdown("#### 📅 Haftalık Ders İçeriği")
                                contents = display_data.get('CourseContent', [])
                                if contents and isinstance(contents, list):
                                    for i, item in enumerate(contents, 1): st.markdown(f"**{i}. Hafta:** {item}")
                                else: st.warning("İçerik girilmemiş.")
                            with col_outcomes:
                                st.markdown("#### 🧠 Öğrenme Çıktıları")
                                outcomes = display_data.get('LearningOutcomes', [])
                                if outcomes and isinstance(outcomes, list):
                                    for item in outcomes: st.markdown(f"- {item}")
                                else: st.warning("Çıktı girilmemiş.")

                        with d_tab3:
                            st.markdown("#### ⚖️ AKTS / İş Yükü Tablosu")
                            workload = display_data.get('WorkloadDetails', [])
                            if workload:
                                df_work = pd.DataFrame(workload)
                                df_work.columns = ["Etkinlik", "Sayısı", "Süresi (Saat)", "Toplam Yük"]
                                st.dataframe(df_work, use_container_width=True, hide_index=True)
                            
                            st.divider()
                            st.markdown("#### 📊 Değerlendirme Yöntemleri")
                            evaluation = display_data.get('EvaluationMethods', [])
                            if evaluation:
                                df_eval = pd.DataFrame(evaluation)
                                df_eval.columns = ["Yöntem", "Sayısı", "Katkı Yüzdesi (%)"]
                                st.dataframe(df_eval, use_container_width=True, hide_index=True)

                        with d_tab4:
                            st.markdown("#### 🌟 Program Yeterliliklerine Katkısı")
                            contrib = display_data.get('ProgramContribute', [])
                            if contrib:
                                for item in contrib:
                                    level = int(item.get('level', 0))
                                    stars = "⭐" * level + "☆" * (5 - level)
                                    st.markdown(f"**{stars}** ({level}) - {item.get('text', '')}")
                            else: st.info("Veri yok.")

    # ----------------------------------------------------
    # 3. SEKME: DBP Linkinden Çekme
    # ----------------------------------------------------
    with tab_fetch:
        st.subheader("🌐 Harici DBP Linkinden Veri Çek")
        st.info("⚠️ Sadece **Erciyes Üni. DBP** linkleri desteklenmektedir.")
        
        dbp_link = st.text_input("DBP Linkini Buraya Yapıştırın", key="dbp_link_input")
        
        if st.button("Veriyi Çek ve Önizle", type="primary", key="btn_fetch_dbp"):
            if not dbp_link:
                st.error("Lütfen geçerli bir DBP linki girin.")
            elif 'dbp_fetcher' not in globals():
                st.error("DBPFetcher nesnesi bulunamadı.")
            else:
                try:
                    with st.spinner("Veri çekiliyor..."):
                        fetched_data = dbp_fetcher.fetch_course_data(dbp_link)

                    if "error" in fetched_data:
                        st.error(fetched_data["error"])
                    else:
                        st.session_state['fetched_dbp_data'] = fetched_data
                        st.session_state['dbp_link_used'] = dbp_link
                        st.success(f"✅ Veriler çekildi: {fetched_data['CourseCode']}")
                        st.rerun()
                except Exception as e:
                    st.error(f"Hata: {e}")

        if 'fetched_dbp_data' in st.session_state:
            data = st.session_state['fetched_dbp_data']
            st.markdown("---")
            st.subheader(f"📝 Önizleme: {data.get('CourseCode')} - {data.get('CourseName')}")
            
            with st.form("save_fetched_dbp_form"):
                c_save1, c_save2 = st.columns(2)
                final_code = c_save1.text_input("Ders Kodu", value=data.get('CourseCode', 'KOD').upper(), key="final_code_dbp")
                final_name = c_save2.text_input("Ders Adı", value=data.get('CourseName', 'Ad'), key="final_name_dbp")
                
                st.info("Fakülte ve Program bilgisi otomatik çekimde gelmeyebilir, sonradan düzenleyebilirsiniz.")

                if st.form_submit_button("💾 DBP'yi Veritabanına Kaydet", type="primary"):
                    course_data = {
                        'CourseCode': final_code, 'CourseName': final_name, 
                        'Faculty': "Otomatik Çekildi", 'Program': "Otomatik Çekildi", # Varsayılan değer
                        'CourseLevel': data.get('CourseLevel'), 'CourseType': data.get('CourseType'),
                        'CoursePeriod': data.get('CoursePeriod'), 'LocalCredit': data.get('LocalCredit'),
                        'AKTSCredit': data.get('AKTSCredit'), 'CourseLanguage': data.get('CourseLanguage'),
                        'Coordinator': data.get('Coordinator'), 'CourseGoal': data.get('CourseGoal'),
                        'CourseDesc': data.get('CourseDesc'), 'Prerequisites': data.get('Prerequisites'),
                        'Resources': data.get('Resources'),
                        'CourseContent': data.get('CourseContent', []),
                        'LearningOutcomes': data.get('LearningOutcomes', []),
                        'ProgramContribute': data.get('ProgramContribute', []),
                        'WorkloadDetails': data.get('WorkloadDetails', []),
                        'EvaluationMethods': data.get('EvaluationMethods', []),
                    }
                    
                    if db.create_detailed_course(course_data, user['Username']):
                        st.toast(f"DBP '{final_code}' eklendi!", icon="✅")
                        del st.session_state['fetched_dbp_data']
                        st.session_state['course_rerun'] = True
                    else:
                        st.error("Kayıt hatası.")

def question_bank_page():
    user = st.session_state['user']
    st.title("🗃️ Soru Bankası Yönetimi")
    
    all_q = db.get_questions(user)
    if not all_q:
        st.warning("📭 Soru bankası boş. 'Soru Ekle' menüsünden içerik ekleyin.")
        return

    all_courses = db.get_courses(user)
    
    df_questions = pd.DataFrame(all_q)
    df_courses = pd.DataFrame(all_courses)
    
    if not df_courses.empty:
        df_courses_slim = df_courses[['CourseCode', 'Faculty', 'Program']]
        df = pd.merge(df_questions, df_courses_slim, on='CourseCode', how='left')
        
        df['Faculty'] = df['Faculty'].fillna('Tanımsız')
        df['Program'] = df['Program'].fillna('Tanımsız')
    else:
        df = df_questions
        df['Faculty'] = 'Veri Yok'
        df['Program'] = 'Veri Yok'

    with st.expander("🔎 Filtreleme ve Arama", expanded=True):

        tum_fakulteler = sorted([str(x) for x in df['Faculty'].dropna().unique()])
        
        c_f1, c_f2 = st.columns(2)
        sel_faculty = c_f1.multiselect("Fakülte", tum_fakulteler)

        if sel_faculty:
            fakulteye_gore_df = df[df['Faculty'].isin(sel_faculty)]
            uygun_bolumler = sorted([str(x) for x in fakulteye_gore_df['Program'].dropna().unique()])
        else:
            uygun_bolumler = sorted([str(x) for x in df['Program'].dropna().unique()])

        sel_program = c_f2.multiselect("Bölüm/Program", uygun_bolumler)

        temp_df = df.copy()
        
        if sel_faculty:
            temp_df = temp_df[temp_df['Faculty'].isin(sel_faculty)]
        
        if sel_program:
            temp_df = temp_df[temp_df['Program'].isin(sel_program)]
        
        uygun_dersler = sorted(temp_df['CourseCode'].unique())
        uygun_konular = sorted(temp_df['TopicArea'].unique())

        c1, c2 = st.columns(2)
        sel_course = c1.multiselect("Ders Kodları", uygun_dersler)
        sel_topic = c2.multiselect("Konu Alanları", uygun_konular)
        
        c3, c4, c5 = st.columns([1, 1, 2])
        type_map = {"MC": "Çoktan Seçmeli", "TF": "Doğru/Yanlış", "CL": "Klasik"}
        sel_type = c3.multiselect("Soru Tipi", ["MC", "TF", "CL"], format_func=lambda x: type_map.get(x, x))
        sel_diff = c4.multiselect("Zorluk Seviyesi", [1, 2, 3], format_func=lambda x: f"Seviye {x}")
        search_txt = c5.text_input("Metin Ara", placeholder="Soru veya Konu içinde ara...")

    if sel_course: df = df[df['CourseCode'].isin(sel_course)]
    if sel_topic: df = df[df['TopicArea'].isin(sel_topic)]
    if sel_faculty: df = df[df['Faculty'].isin(sel_faculty)]
    if sel_program: df = df[df['Program'].isin(sel_program)]
    if sel_type: df = df[df['QuestionType'].isin(sel_type)]
    if sel_diff: df = df[df['Complexity'].isin(sel_diff)]
    if search_txt: 
        df = df[
            df['QuestionText'].str.contains(search_txt, case=False, na=False) |
            df['TopicArea'].str.contains(search_txt, case=False, na=False)
        ]
    
    col_res, col_btn = st.columns([3, 1])
    col_res.markdown(f"**Sonuç:** **{len(df)}** kayıt bulundu.")
    
    excel_data = BytesIO()
    with pd.ExcelWriter(excel_data, engine='xlsxwriter') as writer:
        df_export = df.drop(columns=['Options']).copy() 
        df_export.to_excel(writer, index=False, sheet_name='Sorular')
    
    col_btn.download_button(
        label="📥 Excel Olarak İndir",
        data=excel_data.getvalue(),
        file_name=f"sorular_{datetime.now().strftime('%Y%m%d')}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        use_container_width=True
    )

    df_editor = df.copy()
    df_editor.insert(0, "Seç", False)
    cols_to_display = ['Seç', 'QuestionID', 'CourseCode', 'Faculty', 'Program', 'TopicArea', 'QuestionType', 'Score', 'Complexity', 'QuestionText', 'UsageCount', 'LastEditedBy', 'LastEditedAt']
    cols_to_disable = [col for col in cols_to_display if col != "Seç"]

    edited_df = st.data_editor(
        df_editor[cols_to_display],
        column_config={
            "Seç": st.column_config.CheckboxColumn(required=True),
            "QuestionID": st.column_config.NumberColumn("ID", width="small"),
            "CourseCode": st.column_config.TextColumn("Ders Kodu"),
            "Faculty": st.column_config.TextColumn("Fakülte", width="medium"),
            "Program": st.column_config.TextColumn("Bölüm", width="medium"),
            "QuestionText": st.column_config.TextColumn("Soru Metni", width="large"),
            "QuestionType": st.column_config.TextColumn("Tip"),
            "Complexity": st.column_config.NumberColumn("Zorluk"),
            "UsageCount": st.column_config.NumberColumn("Kullanım Sayısı", help="Sınavda kullanılma adedi"),
            "LastEditedBy": st.column_config.TextColumn("Son Düzenleyen"),
            "LastEditedAt": st.column_config.DatetimeColumn("Düzenleme Tarihi"),
        },
        use_container_width=True,
        hide_index=True,
        height=400,
        disabled=cols_to_disable  
    )
   
    selected_rows = edited_df[edited_df['Seç']]
    
    if not selected_rows.empty:
        st.divider()
        st.subheader("🛠️ İşlemler")
        col_act1, col_act2 = st.columns([1, 3])
        
        with col_act1:
            if st.button(f"🗑️ Seçili {len(selected_rows)} Soruyu Sil", type="primary", use_container_width=True):
                ids_to_del = selected_rows['QuestionID'].tolist()
                db.bulk_delete_questions(ids_to_del)
                st.toast(f"{len(ids_to_del)} soru silindi!", icon="✅")
                time.sleep(1)
                st.rerun()
        
        if len(selected_rows) == 1:
            q_row = selected_rows.iloc[0]
            q_id = int(q_row['QuestionID'])
            full_data = next((q for q in all_q if q['QuestionID'] == q_id), None)
            
            if full_data:
                with col_act2:
                    with st.container(border=True):
                        show_question_edit_form(q_id, full_data)
                        
                        with st.expander("🧠 AI Pedagojik Analiz (Bloom)", expanded=False):
                            st.info("Bu özellik, seçili sorunun Bloom Taksonomisi'ne göre seviyesini ölçer ve geliştirme önerisi sunar.")
                            
                            if st.button("Soruyu Analiz Et", key=f"analyze_bloom_{q_id}"):
                                with st.spinner("Gemini soruyu pedagojik olarak inceliyor..."):
                                    text_to_analyze = full_data['QuestionText'] 
                                    
                                    analysis_result = AIGenerator.analyze_question_bloom(text_to_analyze)
                                    
                                    if isinstance(analysis_result, dict) and "bloom_level" in analysis_result:
                                        st.markdown(f"### 📊 Sonuç: {analysis_result['bloom_level']}")
                                        st.success(f"**Gerekçe:** {analysis_result.get('reason', '-')}")
                                        st.warning(f"**💡 Geliştirme Önerisi:** {analysis_result.get('improvement_suggestion', '-')}")
                                    elif "error" in analysis_result:
                                        st.error(f"Hata oluştu: {analysis_result['error']}")
                                    else:
                                        st.error("AI'dan beklenen formatta yanıt alınamadı.")

        elif len(selected_rows) > 1:
            with col_act2:
                st.info("Tek bir soruyu düzenlemek için lütfen sadece bir satır seçin.")

def add_question_page():
    st.title("➕ Soru Ekleme Merkezi")

    t1, t2, t3 = st.tabs(["✍️ Manuel Ekleme", "📂 Excel Yükleme", "🤖 AI Soru Asistanı"])
    
    with t1:
        available_courses = db.get_courses(st.session_state['user'])
        
        if not available_courses:
            st.warning("⚠️ Soru ekleyebilmek için önce 'Dersler' menüsünden bir ders oluşturmalısınız.")
            st.stop()

        course_options = {
            f"{c['CourseCode']} - {c['CourseName']} ({c.get('Faculty', 'Genel')})": c['CourseCode'] 
            for c in available_courses
        }
        c1, c2, c3, c4 = st.columns(4)
        
        selected_label = c1.selectbox("Ders Seç", list(course_options.keys()))
        qc = course_options[selected_label]
        
        qt = c2.text_input("Konu", "Genel")
        qdiff = c3.slider("Zorluk Seviyesi", 1, 3, 2)
        qtype = c4.selectbox("Soru Tipi", ["MC", "TF", "CL"], format_func=lambda x: {"MC": "Çoktan Seçmeli", "TF": "Doğru/Yanlış", "CL": "Klasik"}.get(x))

        selected_course_data = next((item for item in available_courses if item["CourseCode"] == qc), None)
        if selected_course_data:
            st.caption(f"📍 **Detay:** {selected_course_data.get('Faculty', '-')} | {selected_course_data.get('Program', '-')}")

        with st.form("add_manual_form", clear_on_submit=False):
            qtext = st.text_area("Soru Metni", height=100, placeholder="Soru metnini buraya giriniz...")
            
            opts = {}
            if qtype == "MC":
                st.write("Seçenekler:")
                oc1, oc2 = st.columns(2)
                opts['A'] = oc1.text_input("A)", placeholder="Seçenek A")
                opts['B'] = oc2.text_input("B)", placeholder="Seçenek B")
                opts['C'] = oc1.text_input("C)", placeholder="Seçenek C")
                opts['D'] = oc2.text_input("D)", placeholder="Seçenek D")
                opts['E'] = oc1.text_input("E)", placeholder="Seçenek E (Opsiyonel)")
                opts = {k:v for k,v in opts.items() if v.strip()}
            
            fc1, fc2 = st.columns(2)
            qans = fc1.text_input("Doğru Cevap (Örn: A, B veya Doğru/Yanlış)", placeholder="MC için A/B/C/D, TF için Doğru/Yanlış")
            qscore = fc2.number_input("Varsayılan Puan", 1, 100, 10)
            
            submitted = st.form_submit_button("Soru Ekle", type="primary")
            
            if submitted:
                if qc and qtext:
                    existing_qs = db.get_questions(st.session_state['user'], course_code=qc)
                    similars = check_similarity(qtext, existing_qs)
                    
                    if similars and not st.session_state.get('force_add_confirm', False):
                        st.warning("⚠️ Bu soruya çok benzeyen kayıtlar bulundu:")
                        for s in similars:
                            st.write(f"- (ID: {s['ID']}, Benzerlik: %{s['Ratio']}) {s['Text'][:60]}...")
                        
                        st.info("Eğer yine de eklemek istiyorsanız, aşağıdaki kutucuğu işaretleyip tekrar 'Soru Ekle'ye basın.")
                        st.checkbox("Benzerliğe rağmen kaydet", key="force_add_confirm")
                    
                    else:
                        if db.add_question({
                            'CourseCode': qc, 'TopicArea': qt, 'Complexity': qdiff,
                            'QuestionType': qtype, 'Score': qscore, 'QuestionText': qtext,
                            'Options': opts, 'CorrectAnswer': qans, 'CreatedBy': st.session_state['user']['Username']
                        }):
                           st.toast("Soru başarıyla eklendi!", icon="🎉")
                           if 'force_add_confirm' in st.session_state: del st.session_state['force_add_confirm']
                else:
                    st.error("Ders Kodu ve Soru Metni zorunludur.")

    # --- 2. EXCEL YÜKLEME SEKMESİ ---
    with t2:
        st.info("Toplu soru yüklemek için şablonu kullanın.")
       
        demo_data = pd.DataFrame([{"CourseCode": "MAT101", "QuestionText": "Örnek", "QuestionType": "CL"}])
        buffer = BytesIO()
        with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer:
            demo_data.to_excel(writer, index=False)
        st.download_button("📥 Örnek Excel Şablonunu İndir", data=buffer.getvalue(), file_name="soru_yukleme_sablonu.xlsx")
        
        up_file = st.file_uploader("Excel Dosyası Yükle (.xlsx)", type=['xlsx'])
        if up_file and st.button("Soruları İçeri Aktar", type="primary"):
            try:
                df_up = pd.read_excel(up_file)
                df_up = df_up.fillna('')
                success_count = 0
                for index, row in df_up.iterrows():
                     if db.add_question({
                        'CourseCode': str(row.get('CourseCode')),
                        'TopicArea': str(row.get('TopicArea', 'Genel')),
                        'Complexity': int(row.get('Complexity', 2)), 
                        'QuestionType': str(row.get('QuestionType', 'CL')),
                        'Score': float(row.get('Score', 10)),
                        'QuestionText': str(row.get('QuestionText')),
                        'Options': {},
                        'CorrectAnswer': str(row.get('CorrectAnswer', '')),
                        'CreatedBy': st.session_state['user']['Username']
                    }):
                        success_count += 1
                st.success(f"{success_count} soru eklendi.")
            except Exception as e:
                st.error(f"Hata: {e}")

    with t3:
        st.markdown("### 🤖 Yapay Zeka ile Soru Üret")
        st.info("⚠️ JSON Modu kullanıldığı için çıktı formatı daha güvenilirdir.")
        
        user_courses = db.get_courses(st.session_state['user'])
        course_codes = [c['CourseCode'] for c in user_courses]
        if not course_codes: course_codes = ["GENEL"]

        c_ai1, c_ai2, c_ai3, c_ai4 = st.columns([1, 1, 1, 1])
        
        ai_provider = c_ai1.radio("AI Sağlayıcı", ["Google Gemini"], horizontal=True) 
        provider_code = "google"
        
        target_course_code = c_ai2.selectbox("Hedef Ders", course_codes, key="ai_target_course")
        
        num_q = c_ai3.slider("Soru Sayısı", 1, 10, 3)
        ai_model = c_ai4.selectbox("Model", ["gemini-2.5-flash", "gemini-2.5-pro"], key="ai_model_select").split(" ")[0]

        st.write("---")
        st.markdown("**🔍 Kaynak Seçimi:**")
        
        col_src1, col_src2 = st.columns([1, 2])
        use_dbp = col_src1.checkbox("📚 Kayıtlı DBP/Ders İçeriğini Kullan", value=True, help="Seçilen dersin veritabanındaki haftalık konularını ve çıktılarını kullanır.")
        
        dbp_context_text = ""
        if use_dbp:
            dbp_context_text = db.get_course_context_for_ai(target_course_code)
            if dbp_context_text:
                st.success(f"✅ '{target_course_code}' için içerik veritabanından çekildi.")
                with st.expander("Çekilen İçeriği Gör"):
                    st.text(dbp_context_text)
            else:
                st.warning(f"⚠️ '{target_course_code}' için veritabanında detaylı içerik bulunamadı.")

        tab_ai_file, tab_ai_text = st.tabs(["📄 Ek Dosya Yükle (Opsiyonel)", "📝 Ek Metin (Opsiyonel)"])
        
        user_uploaded_text = ""
        with tab_ai_file:
            uploaded_doc = st.file_uploader("Ek Ders Notu (PDF/Docx)", type=['pdf', 'docx', 'txt'])
            if uploaded_doc:
                user_uploaded_text = AIGenerator.extract_text_from_file(uploaded_doc)
        
        with tab_ai_text:
            pasted_text = st.text_area("Ek Metin Yapıştır", height=100)
            if pasted_text:
                user_uploaded_text += "\n" + pasted_text
        
        if st.button("🚀 Soruları Oluştur", type="primary"):
            final_prompt_text = ""
            
            if use_dbp and dbp_context_text:
                final_prompt_text += f"--- DERSİN RESMİ MÜFREDATI ---\n{dbp_context_text}\n"
            
            if user_uploaded_text:
                final_prompt_text += f"\n--- EK DERS NOTLARI ---\n{user_uploaded_text}\n"

            if not AIGenerator.get_api_key(provider_code):
                st.error("API Anahtarı eksik.")
            elif len(final_prompt_text) < 20: # En azından biraz metin olmalı
                st.warning("⚠️ Soru üretmek için yeterli içerik yok. Lütfen 'DBP Kullan'ı seçin veya bir dosya yükleyin.")
            else:
                with st.spinner(f"{ai_provider} müfredatı ve notları analiz ediyor..."):
                    qs = AIGenerator.generate_from_text(final_prompt_text, num_q, provider_code, ai_model, target_course=target_course_code)
                    if qs:
                        st.session_state['ai_questions'] = qs
                        st.success(f"✅ {len(qs)} adet soru oluşturuldu.")

        if 'ai_questions' in st.session_state and st.session_state['ai_questions']:
            st.divider()
            st.markdown("#### 📝 Üretilen Soru Taslakları")
            
            for idx, q in enumerate(st.session_state['ai_questions']):
                with st.expander(f"Soru {idx+1}: {q.get('QuestionText', '')[:60]}...", expanded=False):
                    col_q1, col_q2 = st.columns([3, 1])
                    with col_q1:
                        st.markdown(f"**Soru:** {q.get('QuestionText')}")
                        if q.get('Options'):
                            st.caption("Seçenekler:")
                            st.json(q.get('Options'))
                    with col_q2:
                        st.info(f"Cevap: {q.get('CorrectAnswer')}")
                        st.caption(f"Zorluk: {q.get('Complexity')} | Puan: {q.get('Score')}")
                        st.caption(f"Ders Kodu: {q.get('CourseCode')} | Konu: {q.get('TopicArea')}")
                    
                    save_q = q.copy()
                    save_q['Score'] = float(save_q.get('Score', 10))
                    save_q['Complexity'] = int(save_q.get('Complexity', 2))
                    
                    if st.button(f"💾 Veritabanına Ekle (Soru {idx+1})", key=f"btn_ai_save_{idx}"):
                        if db.add_question(save_q):
                           st.toast(f"Soru {idx+1} eklendi!", icon="✅")

def shuffle_question_options(questions_list):
    """
    Çoktan Seçmeli soruların şıklarını ve doğru cevap anahtarını tamamen karıştırır.
    (Anahtarlar ve metinler birlikte karışır)
    """
    shuffled_questions = []
    
    standard_keys = ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H'] 
    
    for q_orig in questions_list:
        q = q_orig.copy() 
        
        if q['QuestionType'] == 'MC':
            raw_opts = q.get('Options')
            opts = raw_opts if isinstance(raw_opts, dict) else (json.loads(raw_opts) if isinstance(raw_opts, str) else {})
            
            valid_opts = {k: v for k, v in opts.items() if v and str(v).strip()}
            
            if valid_opts and len(valid_opts) > 1:
                old_correct_key = q.get('CorrectAnswer', '').strip().upper()
                correct_text = valid_opts.get(old_correct_key)
                
                if correct_text:
                    values = list(valid_opts.values())
                    random.shuffle(values)
                    
                    used_key_count = len(values)
                    new_keys = standard_keys[:used_key_count]
                    
                    new_opts = dict(zip(new_keys, values))
                    
                    new_correct_key = None
                    for k, v in new_opts.items():
                        if v == correct_text:
                            new_correct_key = k
                            break
                    
                    q['Options'] = new_opts
                    q['CorrectAnswer'] = new_correct_key if new_correct_key else old_correct_key
                    
        shuffled_questions.append(q)
    
    return shuffled_questions

def load_draft_exam(exam_id):
    """Veritabanından taslak sınavı yükler ve session state'i günceller."""
    exam_data_row = db.get_single_exam(exam_id)
    if not exam_data_row:
        st.error("Sınav verisi bulunamadı.")
        return False
        
    exam_meta = {
        'course': exam_data_row['CourseCode'], 
        'title': exam_data_row['Title'], 
        'score': exam_data_row['TotalScore'], 
        'method': 'Manual', 
        'creator': exam_data_row['CreatedBy'], 
        'groups': "Sadece A", 
        'classical_lines': 5 
    }
    
    try:
        qs = json.loads(exam_data_row['ExamData'])
        temp_scores = {q['QuestionID']: float(q['Score']) for q in qs if 'QuestionID' in q}
    except Exception as e:
        st.error(f"Taslak veri çözümleme hatası: {e}")
        return False

    st.session_state['exam_meta'] = exam_meta
    st.session_state['selected_questions'] = qs
    st.session_state['temp_scores'] = temp_scores
    st.session_state['exam_stage'] = 'preview' 
    st.session_state['override_score_check'] = False 
    return True

def exam_create_page():
    st.title("⚙️ Sınav Sihirbazı v5.2")
    user = st.session_state['user']
    
    if 'exam_stage' not in st.session_state:
        st.session_state['exam_stage'] = 'setup'
        st.session_state['selected_questions'] = []
    
    if 'temp_scores' not in st.session_state:
        st.session_state['temp_scores'] = {} 
        
    if 'override_score_check' not in st.session_state:
        st.session_state['override_score_check'] = False

    if st.session_state['exam_stage'] == 'setup':
        st.info("Adım 1/3: Sınav Ayarları")

        draft_exams = db.get_exams(user, status='Draft', is_archived=0)
        draft_options = [{"ID": ex['ExamID'], "Title": f"{ex['Title']} ({ex['CourseCode']})"} for ex in draft_exams]
        draft_titles = ["Yeni Sınav Oluştur"] + [opt['Title'] for opt in draft_options]
        
        sel_draft_title = st.selectbox("Taslak Sınav Yükle", draft_titles, index=0)
        
        if sel_draft_title != "Yeni Sınav Oluştur":
            draft_id = next(opt['ID'] for opt in draft_options if opt['Title'] == sel_draft_title)
            if st.button(f"Taslak Sınavı Yükle: {sel_draft_title.split('(')[0]}", type="secondary"):
                if load_draft_exam(draft_id):
                    st.success("Taslak sınav başarıyla yüklendi. Önizleme adımına geçebilirsiniz.")
                    st.rerun()
                else:
                    st.error("Taslak sınav yüklenirken bir hata oluştu.")
        
        st.divider()
        st.markdown("#### Yeni Sınav Bilgileri")
        
        pool_all = db.get_questions(user) 
        courses = sorted(list(set(q['CourseCode'] for q in pool_all)))
        
        c1, c2 = st.columns(2)
        sel_course = c1.selectbox("Ders", courses) if courses else st.selectbox("Ders", ["Veri Yok"])
        sel_title = c2.text_input("Başlık", f"{sel_course} Final" if sel_course != "Veri Yok" else "Yeni Sınav Başlığı")
        
        c3, c4, c5, c6 = st.columns(4)
        sel_score = c3.number_input("Hedef Toplam Puan", value=100, min_value=1) 
        method = c4.radio("Yöntem", ["🎲 Rastgele", "✅ Manuel"], horizontal=True)
        groups = c5.selectbox("Kitapçıklar", ["Sadece A", "A ve B", "A, B, C, D"])
        classical_lines = c6.slider("Klasik Soru Cevap Satırı", 1, 15, 5) 
        
        if st.button("İleri ➡️", type="primary", use_container_width=True) and sel_course != "Veri Yok":
            st.session_state['exam_meta'] = {
                'course': sel_course, 'title': sel_title, 'score': sel_score, 
                'method': method, 'creator': user['Username'], 'groups': groups,
                'classical_lines': classical_lines
            }
            st.session_state['exam_stage'] = 'selection'
            st.session_state['temp_scores'] = {} 
            st.session_state['selected_questions'] = []
            st.session_state['override_score_check'] = False
            st.rerun()
        elif sel_course == "Veri Yok":
            st.warning("Lütfen soru bankasına soru ekledikten sonra devam edin.")

    elif st.session_state['exam_stage'] == 'selection':
        meta = st.session_state['exam_meta']
        st.subheader(f"Adım 2/3: Soru Seçimi ({meta['course']} - {meta['title']})")
        
        if st.button("⬅️ Geri Dön (Ayarlar)", key="back_to_setup"):
            st.session_state['exam_stage'] = 'setup'
            st.rerun()
            
        st.divider()
            
        pool = db.get_questions(user, meta['course'])
        if not pool:
            st.warning(f"'{meta['course']}' dersi için soru havuzunda hiç soru bulunmamaktadır.")
            return

        tab_random, tab_manual = st.tabs(["🎲 Tip ve Zorluk Bazlı (Rastgele)", "✍️ Listeden Seç (Manuel)"])
        
        with tab_random:
            st.info("Sistemin belirlediğiniz kriterlerdeki havuzlardan rastgele soru seçmesi için adetleri girin.")
            
            st.markdown("#### 1. Konu Alanı Kısıtlaması (Opsiyonel)")
            topics_in_course = sorted(list(set(q['TopicArea'] for q in pool)))
            sel_random_topics = st.multiselect("Sadece Şu Konu Alanlarından Seç", topics_in_course, key="rnd_topics")
            
            random_pool = [q for q in pool if (not sel_random_topics or q['TopicArea'] in sel_random_topics)]

            st.markdown("#### 2. Soru Tipi ve Zorluk Adet Seçimi")
            
            pools = {
                'MC_1': [q for q in random_pool if q['QuestionType'] == 'MC' and q['Complexity'] == 1],
                'MC_2': [q for q in random_pool if q['QuestionType'] == 'MC' and q['Complexity'] == 2],
                'MC_3': [q for q in random_pool if q['QuestionType'] == 'MC' and q['Complexity'] == 3],
                'TF_1': [q for q in random_pool if q['QuestionType'] == 'TF' and q['Complexity'] == 1],
                'TF_2': [q for q in random_pool if q['QuestionType'] == 'TF' and q['Complexity'] == 2],
                'TF_3': [q for q in random_pool if q['QuestionType'] == 'TF' and q['Complexity'] == 3],
                'CL_1': [q for q in random_pool if q['QuestionType'] == 'CL' and q['Complexity'] == 1],
                'CL_2': [q for q in random_pool if q['QuestionType'] == 'CL' and q['Complexity'] == 2],
                'CL_3': [q for q in random_pool if q['QuestionType'] == 'CL' and q['Complexity'] == 3],
            }
            
            requested_qs = {}
            total_req = 0
            
            st.markdown("##### 🟦 Çoktan Seçmeli (MC)")
            c_mc1, c_mc2, c_mc3 = st.columns(3)
            req_mc1 = c_mc1.number_input(f"Kolay (Mevcut: {len(pools['MC_1'])})", 0, len(pools['MC_1']), 0, key="req_mc1")
            req_mc2 = c_mc2.number_input(f"Orta (Mevcut: {len(pools['MC_2'])})", 0, len(pools['MC_2']), 0, key="req_mc2")
            req_mc3 = c_mc3.number_input(f"Zor (Mevcut: {len(pools['MC_3'])})", 0, len(pools['MC_3']), 0, key="req_mc3")
            total_req += req_mc1 + req_mc2 + req_mc3
            requested_qs.update({'MC_1': req_mc1, 'MC_2': req_mc2, 'MC_3': req_mc3})

            st.markdown("##### 🟩 Doğru/Yanlış (TF)")
            c_tf1, c_tf2, c_tf3 = st.columns(3)
            req_tf1 = c_tf1.number_input(f"Kolay (Mevcut: {len(pools['TF_1'])})", 0, len(pools['TF_1']), 0, key="req_tf1")
            req_tf2 = c_tf2.number_input(f"Orta (Mevcut: {len(pools['TF_2'])})", 0, len(pools['TF_2']), 0, key="req_tf2")
            req_tf3 = c_tf3.number_input(f"Zor (Mevcut: {len(pools['TF_3'])})", 0, len(pools['TF_3']), 0, key="req_tf3")
            total_req += req_tf1 + req_tf2 + req_tf3
            requested_qs.update({'TF_1': req_tf1, 'TF_2': req_tf2, 'TF_3': req_tf3})

            st.markdown("##### 🟨 Klasik (CL)")
            c_cl1, c_cl2, c_cl3 = st.columns(3)
            req_cl1 = c_cl1.number_input(f"Kolay (Mevcut: {len(pools['CL_1'])})", 0, len(pools['CL_1']), 0, key="req_cl1")
            req_cl2 = c_cl2.number_input(f"Orta (Mevcut: {len(pools['CL_2'])})", 0, len(pools['CL_2']), 0, key="req_cl2")
            req_cl3 = c_cl3.number_input(f"Zor (Mevcut: {len(pools['CL_3'])})", 0, len(pools['CL_3']), 0, key="req_cl3")
            total_req += req_cl1 + req_cl2 + req_cl3
            requested_qs.update({'CL_1': req_cl1, 'CL_2': req_cl2, 'CL_3': req_cl3})

            if total_req > 0:
                score_per_q_rand = meta['score'] / total_req
                st.caption(f"Toplam İstenen Soru: **{total_req}** Soru (Soru Başı ≈ **{score_per_q_rand:.2f}** Puan)")
            
            if st.button("🎲 Rastgele Oluştur & Devam Et", type="primary", use_container_width=True, key="btn_random_create"):
                if total_req == 0:
                    st.warning("En az 1 soru seçmelisiniz.")
                else:
                    selected_qs = []
                    for key, count in requested_qs.items():
                        if count > 0:
                            selected_qs.extend(random.sample(pools[key], count))
                    
                    random.shuffle(selected_qs) 
                    st.session_state['selected_questions'] = selected_qs
                    
                    default_score = meta['score'] / len(selected_qs) if len(selected_qs) > 0 else 0
                    st.session_state['temp_scores'] = {q['QuestionID']: round(default_score, 2) for q in selected_qs}
                    st.session_state['override_score_check'] = False 

                    st.session_state['exam_stage'] = 'preview'
                    st.rerun()

        with tab_manual:
            st.info("Aşağıdaki listeden sınavda sormak istediğiniz soruları işaretleyin. Puanlar bir sonraki adımda ayarlanabilir.")
            
            df = pd.DataFrame(pool)
            df.insert(0, "Seç", False)
            
            edited_df = st.data_editor(
                df[['Seç', 'QuestionID', 'QuestionType', 'Complexity', 'Score', 'QuestionText', 'TopicArea', 'UsageCount']],
                column_config={
                    "Seç": st.column_config.CheckboxColumn(required=True),
                    "QuestionText": st.column_config.TextColumn("Soru", width="large"),
                    "Complexity": st.column_config.NumberColumn("Zorluk"),
                    "Score": st.column_config.NumberColumn("Puan (Varsayılan)"),
                    "UsageCount": st.column_config.NumberColumn("Kullanım Sayısı"),
                },
                disabled=["QuestionID", "QuestionType", "Complexity", "Score", "QuestionText", "TopicArea", "UsageCount"],
                hide_index=True,
                use_container_width=True,
                height=500
            )
            
            manual_selections = edited_df[edited_df["Seç"]]
            count_sel = len(manual_selections)
            
            if count_sel > 0:
                score_per_q_man = meta['score'] / count_sel
                st.write(f"**Seçilen Soru Sayısı:** {count_sel} (Soru Başı ≈ **{score_per_q_man:.2f}** Puan)")
            
            if st.button("✅ Seçilenlerle Oluştur & Devam Et", type="primary", use_container_width=True, key="btn_manual_create"):
                if count_sel == 0:
                    st.warning("Lütfen listeden en az bir soru işaretleyin.")
                else:
                    selected_ids = manual_selections['QuestionID'].tolist()
                    final_selection = [q for q in pool if q['QuestionID'] in selected_ids]
                    
                    st.session_state['selected_questions'] = final_selection

                    default_score = meta['score'] / len(final_selection) if len(final_selection) > 0 else 0
                    st.session_state['temp_scores'] = {q['QuestionID']: round(default_score, 2) for q in final_selection}
                    st.session_state['override_score_check'] = False 

                    st.session_state['exam_stage'] = 'preview'
                    st.rerun()

    elif st.session_state['exam_stage'] == 'preview':
        meta = st.session_state['exam_meta']
        qs = st.session_state['selected_questions']
        
        st.subheader(f"Adım 3/3: Sınav Önizleme ve Puanlama ({len(qs)} Soru)")
        
        c_back, c_save_draft = st.columns([1, 1])
        if c_back.button("⬅️ Geri Dön (Seçim)", key="back_to_selection"):
            st.session_state['exam_stage'] = 'selection'
            st.rerun()
        
        if "score_editor" in st.session_state:
            editor_changes = st.session_state["score_editor"]
            
            if isinstance(editor_changes, dict) and "edited_rows" in editor_changes:
                pass
            elif isinstance(editor_changes, pd.DataFrame):
                for _, row in editor_changes.iterrows():
                    q_id = int(row['Soru ID'])
                    new_val = float(row['Manuel Puan'])
                    st.session_state['temp_scores'][q_id] = new_val
        
        # -------------------------------------------------------------
        
        if c_save_draft.button("💾 Taslak Olarak Kaydet", key="save_as_draft", use_container_width=True):
            final_qs = []
            for q in qs:
                q_id = q['QuestionID']
                q_copy = q.copy()
                q_copy['Score'] = st.session_state['temp_scores'].get(q_id, q['Score']) 
                final_qs.append(q_copy)
            
            meta['score'] = sum(q['Score'] for q in final_qs)

            db.save_exam(meta, final_qs, status='Draft')
            st.toast("Sınav taslak olarak kaydedildi.", icon="💾")
            st.session_state.pop('exam_stage', None)
            st.rerun()

        st.divider()

        default_score_for_init = meta['score'] / len(qs) if len(qs) > 0 else 0
        score_data_init = []
        for i, q in enumerate(qs):
            q_id = q['QuestionID']
            current_score = st.session_state['temp_scores'].get(q_id, round(default_score_for_init, 2))
            
            score_data_init.append({
                "Soru ID": q_id,
                "Soru Metni": q['QuestionText'][:80] + "...",
                "Varsayılan Puan": round(default_score_for_init, 2),
                "Manuel Puan": current_score,
            })
        df_scores_init = pd.DataFrame(score_data_init)
        
        with st.form("score_adjustment_form"):
            st.markdown("#### ⚖️ Puan Ayarlama")
            
            edited_scores_df = st.data_editor(
                df_scores_init,
                column_config={
                    "Soru ID": st.column_config.NumberColumn(disabled=True),
                    "Soru Metni": st.column_config.TextColumn(width="large", disabled=True),
                    "Varsayılan Puan": st.column_config.NumberColumn(disabled=True),
                    "Manuel Puan": st.column_config.NumberColumn(format="%.2f", min_value=0.1)
                },
                hide_index=True,
                use_container_width=True,
                key="score_editor" 
            )
            
            new_total_score = edited_scores_df['Manuel Puan'].sum() if not edited_scores_df.empty else 0.0
            st.markdown(f"**Yeni Toplam Sınav Puanı:** **{new_total_score:.2f}** (Hedef Puan: {meta['score']})")
            
            if st.session_state['override_score_check'] == True and abs(new_total_score - meta['score']) > 0.01:
                submit_label = "⚠️ Farklı Puanla Kaydetmeyi Onayla"
                submit_type = "secondary"
                st.warning(f"⚠️ **Uyarı:** Sınavın hedef puanı **{meta['score']}** iken, manuel puan toplamı **{new_total_score:.2f}** oldu.")
            else:
                 submit_label = "✅ Sınavı Oluştur & Kaydet"
                 submit_type = "primary"
            
            submitted = st.form_submit_button(submit_label, type=submit_type, use_container_width=True)

        if submitted:
            new_temp_scores = {}
            for _, row in edited_scores_df.iterrows():
                q_id = int(row['Soru ID'])
                score = round(float(row['Manuel Puan']), 2)
                new_temp_scores[q_id] = score
            st.session_state['temp_scores'] = new_temp_scores 

            if new_total_score <= 0:
                st.error("Toplam puan 0'dan büyük olmalıdır.")
            
            elif abs(new_total_score - meta['score']) > 0.01 and st.session_state['override_score_check'] == False:
                 st.session_state['override_score_check'] = True
                 st.rerun() 
            
            else:
                final_qs = []
                for q in qs:
                    q_id = q['QuestionID']
                    q_copy = q.copy()
                    q_copy['Score'] = st.session_state['temp_scores'].get(q_id, q_copy['Score']) 
                    final_qs.append(q_copy)
                
                meta['score'] = round(new_total_score, 2)
                
                db.save_exam(meta, final_qs, status='Final')
                st.session_state['final_qs'] = final_qs 
                st.session_state['exam_stage'] = 'finish'
                st.session_state['override_score_check'] = False 
                st.rerun()
            
        with st.expander("### 👁️ Sınav Kağıdı Önizlemesi (A Kitapçığı)", expanded=False):
               
            preview_qs = []
            current_scores = {int(row['Soru ID']): float(row['Manuel Puan']) for _, row in edited_scores_df.iterrows()}
            
            for q in qs:
                q_copy = q.copy()
                q_copy['Score'] = current_scores.get(q_copy['QuestionID'], q_copy['Score'])
                preview_qs.append(q_copy)
            
            preview_container = st.container(border=True, height=500)
            with preview_container:
                st.markdown(f"#### {meta['title']} ({meta['course']}) - A Kitapçığı")
                st.markdown("---")
                for idx, q in enumerate(preview_qs, 1):
                    st.markdown(f"**{idx}.** {q['QuestionText']} **({q['Score']} Puan)**")
                    if q['QuestionType'] == 'MC':
                        opts = q['Options'] if isinstance(q['Options'], dict) else (json.loads(q['Options']) if isinstance(q['Options'], str) else {})
                        if opts:
                             for k, v in sorted(opts.items()): st.write(f"{k}) {v}")
                    st.markdown("---")

    elif st.session_state['exam_stage'] == 'finish':
        st.success("✅ Sınav Başarıyla Hazırlandı!")
        meta = st.session_state['exam_meta']
        base_questions = st.session_state['final_qs'] 
        
        group_list = ["A"]
        if meta['groups'] == "A ve B": group_list = ["A", "B"]
        elif "C" in meta['groups']: group_list = ["A", "B", "C", "D"]
        
        classical_lines = meta.get('classical_lines', 5)

        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zf:
            for grp in group_list:
                if grp == "A":
                    current_qs = [q.copy() for q in base_questions]
                else: 
                    current_qs_shuffled = [q.copy() for q in base_questions]
                    random.shuffle(current_qs_shuffled)
                    current_qs = shuffle_question_options(current_qs_shuffled) 
                    
                pdf = ExamPDFEngine(meta, group_name=grp, classical_lines=classical_lines)
                pdf.generate_content(current_qs)
                zf.writestr(f"SoruKitapcigi_{grp}_{meta['course']}.pdf", pdf.get_pdf_bytes())
                
                pdfk = ExamPDFEngine(meta, is_answer_key=True, group_name=grp, classical_lines=classical_lines)
                pdfk.generate_content(current_qs)
                zf.writestr(f"CevapAnahtari_{grp}_{meta['course']}.pdf", pdfk.get_pdf_bytes())

                docx = ExamDocxEngine(meta, group_name=grp, classical_lines=classical_lines)
                docx.generate(current_qs)
                zf.writestr(f"SoruKitapcigi_{grp}_{meta['course']}.docx", docx.get_docx_bytes())
                
                docxk = ExamDocxEngine(meta, is_answer_key=True, group_name=grp, classical_lines=classical_lines)
                docxk.generate(current_qs)
                zf.writestr(f"CevapAnahtari_{grp}_{meta['course']}.docx", docxk.get_docx_bytes())

        st.download_button(
            "📦 Tüm Sınav Setini İndir (ZIP)",
            data=zip_buffer.getvalue(),
            file_name=f"{meta['course']}_sinav_seti_{datetime.now().strftime('%Y%m%d')}.zip",
            mime="application/zip",
            type="primary",
            use_container_width=True
        )
        if st.button("➕ Yeni Sınav Oluştur", use_container_width=True):
            keys_to_delete = ['exam_stage', 'selected_questions', 'exam_meta', 'final_qs', 'temp_scores', 'override_score_check']
            for key in keys_to_delete:
                if key in st.session_state: del st.session_state[key]
            st.rerun()

def admin_page():
    if st.session_state['user']['Role'] != 'Admin':
        st.error("Bu sayfaya erişim yetkiniz yok.")
        return
        
    st.title("🛡️ Yönetici Paneli")
    
    tab1, tab2, tab3, tab4 = st.tabs(["Kullanıcı Listesi", "Yeni Kullanıcı Ekle", "Veritabanı Yönetimi", "Sistem Logları"])
    
    with tab1:
        st.subheader("Kullanıcı Yönetimi")
        users = db.get_all_users()
        
        st.markdown("### 👤 Mevcut Kullanıcılar")
        
        for u in users:
            with st.container(border=True):
                col_u1, col_u2 = st.columns([4, 1])
                email_display = u['Email'] if u['Email'] else 'E-posta Yok'
                
                with col_u1:
                    st.markdown(f"**{u['Username']}** ({u['FullName']})")
                    st.caption(f"Rol: {u['Role']} | Tema: {u['Theme']} | Email: {email_display}")
                
                with col_u2:
                    if u['Username'] != 'admin':
                        if st.button("🗑️ Sil", key=f"del_{u['Username']}", type="secondary", use_container_width=True):
                            if db.delete_user(u['Username']): 
                                st.toast(f"Kullanıcı silindi.", icon="✅")
                                time.sleep(0.5)
                                st.session_state['rerun_admin'] = True
                    else:
                        st.info("Sistem Yöneticisi")

                with st.expander(f"🔑 {u['Username']} için Şifre İşlemleri"):
                    c_pass1, c_pass2 = st.columns([3, 1])
                    new_pass_input = c_pass1.text_input("Yeni Şifre Girin", type="password", key=f"new_pass_{u['Username']}", placeholder="Yeni şifre...")
                    
                    if c_pass2.button("Güncelle", key=f"btn_rst_{u['Username']}", type="primary"):
                        if db.reset_password(u['Username'], new_pass_input):
                            st.success(f"Şifre başarıyla güncellendi!")
                        else:
                            st.error("Şifre boş olamaz.")

        if st.session_state.get('rerun_admin'):
             st.session_state.pop('rerun_admin')
             st.rerun()

    with tab2:
        st.subheader("➕ Yeni Kullanıcı Oluştur")
        
        ALL_THEMES = [
            "Glassmorphism", "Apple Premium", "Neo Dark", "Material Design", 
            "Neumorphism", "Holographic Dark", "Monochrome", "Corporate Blue", 
            "Retro Terminal", "Tropical Sunset", "Soft Pastel"
        ]
        
        with st.form("new_user"):
            new_u = st.text_input("Kullanıcı Adı")
            new_p = st.text_input("Şifre", type="password")
            new_n = st.text_input("Ad Soyad")
            new_e = st.text_input("E-posta", placeholder="kullanici@domain.com")
            
            c_nu1, c_nu2 = st.columns(2)
            new_r = c_nu1.selectbox("Rol", ["Öğretim Üyesi", "Admin"])
            new_t = c_nu2.selectbox("Tema Seçimi", ALL_THEMES, index=0) 

            if st.form_submit_button("Kullanıcı Oluştur", type="primary"):
                if db.create_user(new_u, new_p, new_n, new_r, new_t, new_e): 
                    st.success(f"Kullanıcı '{new_u}' oluşturuldu!")
                    time.sleep(1)
                    st.rerun() 
                else:
                    st.error("Kullanıcı adı zaten var veya bir veritabanı hatası oluştu.")
                    
    with tab3:
        st.subheader("Veritabanı Yedekleme")
        db_path = DB_FILE
        if os.path.exists(db_path):
            with open(db_path, "rb") as f:
                db_bytes = f.read()
            st.download_button(
                label="📥 SQLite Veritabanını İndir (.sqlite)",
                data=db_bytes,
                file_name=f"ssop_yedek_{datetime.now().strftime('%Y%m%d_%H%M%S')}.sqlite",
                mime="application/octet-stream"
            )
            st.caption("Bu, uygulamanın tüm verilerini (sorular, kullanıcılar, sınavlar) içeren yedek dosyadır.")
            st.markdown("---")
        else:
            st.error("Veritabanı dosyası bulunamadı.")
            
    with tab4:
        st.subheader("Sistem Aksiyon Logları")
        logs = db.get_audit_logs(limit=100) 
        df_logs = pd.DataFrame(logs)
        if not df_logs.empty:
            df_logs['Timestamp'] = df_logs['Timestamp'].apply(lambda x: datetime.strptime(x.split('.')[0], '%Y-%m-%d %H:%M:%S').strftime('%Y-%m-%d %H:%M:%S'))
            st.dataframe(df_logs[['Timestamp', 'Username', 'Action', 'Details']], use_container_width=True, hide_index=True)
        else:
            st.info("Sistem log kaydı bulunmamaktadır.")

def history_page():
    user = st.session_state['user']
    st.title("🗂️ Sınav Arşivi")
    
    tab_active, tab_draft, tab_archived = st.tabs(["Aktif Sınavlar", "Taslak Sınavlar", "Arşivlenmiş Sınavlar"])
    
    with tab_active:
        exams = db.get_exams(user, is_archived=0, status='Final') 
        if not exams:
            st.info("Henüz oluşturulmuş aktif (Final) bir sınav yok.")
            
        for ex in exams:
            col_main, col_action = st.columns([6, 1])
            
            with col_main:
                expander_title = f"📄 {ex['Title']} ({ex['CourseCode']}) | Oluşturulma: {datetime.strptime(ex['CreatedAt'].split('.')[0], '%Y-%m-%d %H:%M:%S').strftime('%d.%m.%Y %H:%M')}"
                
                with st.expander(expander_title):
                    st.markdown(f"**Durum:** {ex['Status']} | **Toplam Puan:** {ex['TotalScore']}")
                    st.markdown(f"**Oluşturan:** {ex['CreatedBy']}")
                    
                    try:
                        q_data = json.loads(ex['ExamData'])
                        st.info(f"Bu sınavda toplam **{len(q_data)}** soru bulunmaktadır.")
                        
                        if st.checkbox("Soruları Listele", key=f"view_q_{ex['ExamID']}"):
                            for i, q in enumerate(q_data, 1):
                                opts_display = q['Options'] if isinstance(q['Options'], dict) else (json.loads(q['Options']) if isinstance(q['Options'], str) else {})
                                opt_str = ', '.join([f"{k}:{v[:20]}" for k, v in opts_display.items()]) if opts_display else "-"
                                st.markdown(f"**{i}. ({q['Score']} P)** **{q['QuestionType']}**: {q['QuestionText'][:100]}...")
                                st.caption(f"Cevap: {q['CorrectAnswer']} | Seçenekler: {opt_str}")

                    except Exception as e:
                        st.error(f"Veri yükleme hatası: {e}")

            with col_action:
                st.write("") 
                if st.button("🗑️ Arşivle", key=f"archive_exam_{ex['ExamID']}", type="secondary", use_container_width=True):
                    db.archive_exam(ex['ExamID'])
                    st.toast(f"Sınav (ID: {ex['ExamID']}) arşivlendi!", icon="✅")
                    time.sleep(1)
                    st.rerun()
                    
    with tab_draft:
        draft_exams = db.get_exams(user, is_archived=0, status='Draft') 
        if not draft_exams:
            st.info("Kaydedilmiş taslak sınav kaydı bulunmamaktadır.")
            return

        for ex in draft_exams:
             col_draft_main, col_draft_action = st.columns([6, 1])
             with col_draft_main:
                 expander_title = f"📝 {ex['Title']} ({ex['CourseCode']}) | Oluşturulma: {datetime.strptime(ex['CreatedAt'].split('.')[0], '%Y-%m-%d %H:%M:%S').strftime('%d.%m.%Y %H:%M')}"
                 with st.expander(expander_title):
                    st.markdown(f"**Durum:** {ex['Status']} | **Toplam Puan:** {ex['TotalScore']}")
                    st.markdown(f"**Oluşturan:** {ex['CreatedBy']}")
             with col_draft_action:
                 st.write("") 
                 if st.button("🛠️ Devam Et", key=f"edit_draft_{ex['ExamID']}", type="primary", use_container_width=True):
                     if load_draft_exam(ex['ExamID']):
                         st.session_state['exam_stage'] = 'preview'
                         st.rerun()
    
    with tab_archived:
        archived_exams = db.get_exams(user, is_archived=1) 
        if not archived_exams:
            st.info("Arşivlenmiş sınav kaydı bulunmamaktadır.")
            return
        
        for ex in archived_exams:
             col_arc_main, col_arc_action = st.columns([6, 1])
             with col_arc_main:
                 st.caption(f"**{ex['Title']}** ({ex['CourseCode']}) | Arşivlendi: {datetime.strptime(ex['CreatedAt'].split('.')[0], '%Y-%m-%d %H:%M:%S').strftime('%d.%m.%Y')}")
             with col_arc_action:
                 st.info("Arşivlendi")

# --------------------------------------------------------------------
# PAGE ROUTER
# --------------------------------------------------------------------
def route_page(selected):
    if selected == "Gösterge Paneli":
        dashboard_page()
    elif selected == "Dersler":
        course_management_page()
    elif selected == "Soru Ekle":
        add_question_page()
    elif selected == "Soru Bankası":
        question_bank_page()
    elif selected == "Sınav Oluştur":
        exam_create_page()
    elif selected == "Arşiv":
        history_page()
    elif selected == "Yönetim":
        admin_page()

def render_sidebar(user, allowed_options, icons):
    """
    Kullanıcının rolüne göre kenar çubuğunu (sidebar) oluşturur,
    tema seçimine izin verir ve Streamlit Option Menu ile menüyü render eder.
    Tema seçimi sadece Admin rolündeki kullanıcılar için görünür.
    """
    role = user.get("Role", "Misafir Kullanıcı") 
    full_name = user.get("FullName", "Kullanıcı Adı")
    email = user.get("Email", "") or ""
    
    ALL_THEMES = [
        "Glassmorphism", "Apple Premium", "Neo Dark", "Material Design", 
        "Neumorphism", "Holographic Dark", "Monochrome", "Corporate Blue", 
        "Retro Terminal", "Tropical Sunset", "Soft Pastel"
    ]
    
    default_theme = user.get("Theme", "Glassmorphism") 
    
    theme = default_theme
    # -------------------------------------------------------------
    # 1. TEMA SEÇİMİ VE GİZLEME MANTIĞI
    # -------------------------------------------------------------
    if role == "Admin":
        st.sidebar.markdown("## ⚙️ Uygulama Ayarları")
        with st.sidebar.expander("🎨 Tema Seçimi", expanded=False):
            theme = st.radio(
                "UI Teması",
                ALL_THEMES,
                index=ALL_THEMES.index(default_theme) if default_theme in ALL_THEMES else 0,
                key="admin_theme_selection",
                help="Uygulamanın görsel temasını seçin. Seçiminiz sadece bu oturum için geçerlidir."
            )
    
    st.sidebar.markdown(
        """
        <style>
            .stSidebar > div:first-child {
                padding-top: 1rem;
            }
            div[data-testid="stRadio"] label {
                margin-bottom: 5px;
            }
        </style>
        """,
        unsafe_allow_html=True
    )
    # ============================ TEMA 1: Glassmorphism ============================
    if theme == "Glassmorphism":
        st.sidebar.markdown(
            """
            <style>
                .glass-card {
                    background: rgba(255, 255, 255, 0.35); 
                    backdrop-filter: blur(9px); 
                    -webkit-backdrop-filter: blur(9px); 
                    border-radius: 18px;
                    padding: 20px;
                    border: 1px solid rgba(255, 255, 255, 0.3); 
                    box-shadow: 0 6px 25px rgba(0, 0, 0, 0.15); 
                    margin-bottom: 15px;
                    transition: all 0.3s ease; 
                }
                .glass-card:hover { box-shadow: 0 8px 30px rgba(0, 0, 0, 0.25); }
            </style>
            """,
            unsafe_allow_html=True
        )
        st.sidebar.markdown(
            f"""
            <div class="glass-card">
                <img src="https://ui-avatars.com/api/?name={full_name}&background=3498db&color=fff&size=110"
                style="border-radius:50%; width:110px; display:block; margin:auto; border: 3px solid #3498db;">
                <h3 style="text-align:center; margin-top:10px; color:#1a242f;">{full_name}</h3>
                <p style="text-align:center; background:#eef7ff; padding:4px; border-radius:10px; color:#2c3e50; font-weight:bold;">{role}</p>
                <p style="text-align:center; font-size:12px; margin-top:5px; color:#5c6d7a;">{email or user.get('Username', 'bilgi')}</p>
            </div>
            """,
            unsafe_allow_html=True
        )
        with st.sidebar:
            st.markdown("### 🧭 Ana Menü")
            selected = option_menu(
                None, allowed_options, icons=icons, key="menu_glass",
                styles={
                    "container": {"background": "rgba(255,255,255,0.15)", "border-radius": "12px", "padding": "10px"},
                    "icon": {"color": "#2980b9", "font-size": "18px"},
                    "nav-link": {"font-size": "16px", "color": "#1a242f", "--hover-color": "rgba(255,255,255,0.5)", "border-radius": "8px", "margin": "5px 0"},
                    "nav-link-selected": {
                        "background": "linear-gradient(90deg, #49a2dd, #6fc3ff)",
                        "color": "white",
                        "font-weight": "bold"
                    }
                }
            )
    # ============================ TEMA 2: Apple Premium ============================
    elif theme == "Apple Premium":
        st.sidebar.markdown(
            """
            <style>
                .apple-card {
                    background: white;
                    border-radius: 22px; 
                    padding: 25px;
                    box-shadow: 0 10px 25px rgba(0, 0, 0, 0.12); 
                    margin-bottom: 20px;
                    border: 1px solid #f0f0f0; 
                    transition: transform 0.2s ease;
                }
                .apple-card:hover { transform: translateY(-2px); }
            </style>
            """,
            unsafe_allow_html=True
        )
        st.sidebar.markdown(
            f"""
            <div class="apple-card">
                <img src="https://ui-avatars.com/api/?name={full_name}&background=2ecc71&color=fff&size=120"
                style="border-radius:50%; width:120px; display:block; margin:auto; box-shadow: 0 0 10px rgba(46, 204, 113, 0.5);">
                <h3 style="text-align:center; color:#333; margin-bottom: 5px;">{full_name}</h3>
                <p style="text-align:center; background:#e8ffe8; padding:5px; border-radius:10px; color:#27ae60; font-weight:600;">{role}</p>
                <p style="text-align:center; font-size:12px; margin-top:5px; color:#7f8c8d;">{email or user.get('Username', 'bilgi')}</p>
            </div>
            """,
            unsafe_allow_html=True
        )
        with st.sidebar:
            st.markdown("### 🚀 Özellikler Menüsü")
            selected = option_menu(
                None, allowed_options, icons=icons, key="menu_apple",
                styles={
                    "icon": {"color": "#27ae60", "font-size": "18px"},
                    "nav-link": {"font-size": "16px", "color": "#333", "border-radius": "10px", "margin": "4px 0"},
                    "nav-link-selected": {"background": "#2ecc71", "color": "white", "font-weight": "bold"},
                }
            )
    # ============================ TEMA 3: Neo Dark ============================
    elif theme == "Neo Dark": 
        st.sidebar.markdown(
            """
            <style>
                .neo-card {
                    background: #0d1117; 
                    padding: 25px;
                    border-radius: 16px;
                    text-align:center;
                    color:white;
                    margin-bottom:18px;
                    border: 1px solid #1f252b; 
                    box-shadow: 0 4px 15px rgba(0, 212, 255, 0.2); 
                }
            </style>
            """,
            unsafe_allow_html=True
        )
        st.sidebar.markdown(
            f"""
            <div class="neo-card">
                <img src="https://ui-avatars.com/api/?name={full_name}&background=00d4ff&color=fff&size=115"
                style="border-radius:50%; width:115px; border: 2px solid #00eaff; box-shadow: 0 0 10px #00eaff;">
                <h3 style="color:#00eaff; margin-bottom: 5px;">{full_name}</h3>
                <p style="background:#001f2b; padding:6px; border-radius:8px; color:#95bcf2; font-weight:500;">{role}</p>
                <p style="font-size:12px; margin-top:5px; color:#8b949e;">{email or user.get('Username', 'bilgi')}</p>
            </div>
            """,
            unsafe_allow_html=True
        )
        with st.sidebar:
            selected = option_menu(
                None, allowed_options, icons=icons, key="menu_dark",
                styles={
                    "container": {"background": "#0d1117", "padding": "10px", "border-radius": "10px"},
                    "icon": {"color": "#00eaff", "font-size": "18px"},
                    "nav-link": {"font-size": "16px", "color": "white", "border-radius": "8px", "margin": "5px 0", "--hover-color": "#161b22"},
                    "nav-link-selected": {"background": "#00eaff", "color": "#0d1117", "font-weight": "bold"},
                }
            )
    # ============================ TEMA 4: Material Design ============================
    elif theme == "Material Design":
        st.sidebar.markdown(
            """
            <style>
                .material-card {
                    background: #f7f9fb; 
                    border-radius: 4px; 
                    padding: 20px;
                    box-shadow: 0 2px 5px rgba(0,0,0,0.2), 0 0 0 1px rgba(0,0,0,0.1); 
                    margin-bottom: 15px;
                }
            </style>
            """,
            unsafe_allow_html=True
        )
        st.sidebar.markdown(
            f"""
            <div class="material-card">
                <img src="https://ui-avatars.com/api/?name={full_name}&background=f44336&color=fff&size=100"
                style="border-radius:50%; width:100px; display:block; margin:auto; border: 4px solid #f44336;">
                <h3 style="text-align:center; color:#212121; margin-top:10px;">{full_name}</h3>
                <p style="text-align:center; background:#e0f7fa; padding:4px; border-radius:4px; color:#00bcd4; font-weight:500;">{role}</p>
                <p style="text-align:center; font-size:12px; margin-top:5px; color:#757575;">{email or user.get('Username', 'bilgi')}</p>
            </div>
            """,
            unsafe_allow_html=True
        )
        with st.sidebar:
            selected = option_menu(
                None, allowed_options, icons=icons, key="menu_material",
                styles={
                    "container": {"background": "#ffffff"},
                    "icon": {"color": "#f44336", "font-size": "18px"},
                    "nav-link": {"font-size": "16px", "color": "#212121", "border-radius": "4px", "margin": "4px 0"},
                    "nav-link-selected": {"background": "#4caf50", "color": "white", "font-weight": "bold"},
                }
            )
    # ============================ TEMA 5: Neumorphism ============================
    elif theme == "Neumorphism":
        st.sidebar.markdown(
            """
            <style>
                .neo-morph-card {
                    background: #ecf0f3; 
                    border-radius: 20px;
                    padding: 25px;
                    box-shadow: 8px 8px 15px #c8c8c8, -8px -8px 15px #ffffff; 
                    margin-bottom: 20px;
                    text-align: center;
                }
            </style>
            """,
            unsafe_allow_html=True
        )
        st.sidebar.markdown(
            f"""
            <div class="neo-morph-card">
                <img src="https://ui-avatars.com/api/?name={full_name}&background=a9ccfb&color=0b5fb0&size=110"
                style="border-radius:50%; width:110px; display:block; margin:auto; box-shadow: 3px 3px 6px #c8c8c8, -3px -3px 6px #ffffff;">
                <h3 style="color:#0b5fb0; margin-top:10px;">{full_name}</h3>
                <p style="text-align:center; color:#0b5fb0; font-weight:600;">{role}</p>
                <p style="text-align:center; font-size:12px; margin-top:5px; color:#607d8b;">{email or user.get('Username', 'bilgi')}</p>
            </div>
            """,
            unsafe_allow_html=True
        )
        with st.sidebar:
            selected = option_menu(
                None, allowed_options, icons=icons, key="menu_neumorphism",
                styles={
                    "container": {"background": "#ecf0f3", "border-radius": "15px", "padding": "10px"},
                    "icon": {"color": "#0b5fb0", "font-size": "18px"},
                    "nav-link": {
                        "font-size": "16px", "color": "#0b5fb0", "border-radius": "15px", "margin": "6px 0",
                        "background": "#ecf0f3", 
                        "box-shadow": "3px 3px 6px #c8c8c8, -3px -3px 6px #ffffff", 
                        "--hover-color": "#e0e3e6"
                    },
                    "nav-link-selected": {
                        "box-shadow": "inset 4px 4px 8px #c8c8c8, inset -4px -4px 8px #ffffff", 
                        "background": "#a9ccfb", 
                        "color": "#0b5fb0",
                        "font-weight": "bold"
                    }
                }
            )
    # ============================ TEMA 6: Holographic Dark ============================
    elif theme == "Holographic Dark":
        st.sidebar.markdown(
            """
            <style>
                .holographic-card {
                    background: linear-gradient(145deg, #1e002e, #000000); 
                    border-radius: 12px;
                    padding: 25px;
                    border: 1px solid #7b00ff; 
                    box-shadow: 0 0 20px rgba(123, 0, 255, 0.7); 
                    margin-bottom: 20px;
                    color: white;
                }
            </style>
            """,
            unsafe_allow_html=True
        )
        with st.sidebar:
            st.sidebar.markdown(
                f"""
                <div class="holographic-card">
                    <img src="https://ui-avatars.com/api/?name={full_name}&background=00ffff&color=1e002e&size=115"
                    style="border-radius:50%; width:115px; border: 3px solid #00ffff; box-shadow: 0 0 10px #00ffff;">
                    <h3 style="color:#00ffff; text-shadow: 0 0 5px #00ffff; margin-bottom: 5px;">{full_name}</h3>
                    <p style="background:#0a001a; padding:6px; border-radius:8px; color:#ff00ff; font-weight:bold;">{role}</p>
                    <p style="font-size:12px; margin-top:5px; color:#ff00ff; text-shadow: 0 0 2px #ff00ff;">{email or user.get('Username', 'bilgi')}</p>
                </div>
                """,
                unsafe_allow_html=True
            )
            selected = option_menu(
                None, allowed_options, icons=icons, key="menu_holographic",
                styles={
                    "container": {"background": "#000000"},
                    "icon": {"color": "#ff00ff", "font-size": "18px"},
                    "nav-link": {
                        "font-size": "16px", "color": "#00ffff", "border-radius": "6px", "margin": "5px 0",
                        "text-shadow": "0 0 2px #00ffff", "--hover-color": "#1a0033"
                    },
                    "nav-link-selected": {
                        "background": "linear-gradient(90deg, #7b00ff, #00ffff)",
                        "color": "#000000",
                        "font-weight": "bold",
                        "border-left": "4px solid #ff00ff"
                    },
                }
            )
    # ============================ YENİ TEMA 7: Monochrome ============================
    elif theme == "Monochrome":
        st.sidebar.markdown(
            """
            <style>
                .mono-card {
                    background: #ffffff;
                    border: 1px solid #1c1c1c; 
                    border-radius: 8px;
                    padding: 20px;
                    box-shadow: 4px 4px 0 #1c1c1c; /* Sert gölge */
                    margin-bottom: 15px;
                    color: #1c1c1c;
                }
            </style>
            """,
            unsafe_allow_html=True
        )    
        st.sidebar.markdown(
            f"""
            <div class="mono-card">
                <img src="https://ui-avatars.com/api/?name={full_name}&background=1c1c1c&color=ffffff&size=100"
                style="border-radius:50%; width:100px; display:block; margin:auto; border: 3px solid #1c1c1c;">
                <h3 style="text-align:center; margin-top:10px; color:#1c1c1c;">{full_name}</h3>
                <p style="text-align:center; background:#e0e0e0; padding:4px; border-radius:4px; color:#1c1c1c; font-weight:bold;">{role}</p>
                <p style="text-align:center; font-size:12px; margin-top:5px; color:#555555;">{email or user.get('Username', 'bilgi')}</p>
            </div>
            """,
            unsafe_allow_html=True
        )
        with st.sidebar:
            selected = option_menu(
                None, allowed_options, icons=icons, key="menu_mono",
                styles={
                    "container": {"background": "#f0f0f0", "border-radius": "6px"},
                    "icon": {"color": "#1c1c1c", "font-size": "18px"},
                    "nav-link": {"font-size": "16px", "color": "#1c1c1c", "border-radius": "4px", "margin": "4px 0", "--hover-color": "#d0d0d0"},
                    "nav-link-selected": {"background": "#1c1c1c", "color": "white", "font-weight": "bold"},
                }
            )
    # ============================ TEMA 8: Corporate Blue ============================
    elif theme == "Corporate Blue":
        st.sidebar.markdown(
            """
            <style>
                .corp-card {
                    background: #ffffff;
                    border-left: 5px solid #0056b3; /* Mavi çizgi */
                    border-radius: 4px;
                    padding: 20px;
                    box-shadow: 0 1px 3px rgba(0,0,0,0.1); 
                    margin-bottom: 15px;
                }
            </style>
            """,
            unsafe_allow_html=True
        )
        st.sidebar.markdown(
            f"""
            <div class="corp-card">
                <img src="https://ui-avatars.com/api/?name={full_name}&background=007bff&color=ffffff&size=100"
                style="border-radius:50%; width:100px; display:block; margin:auto; border: 3px solid #007bff;">
                <h3 style="text-align:center; margin-top:10px; color:#0056b3;">{full_name}</h3>
                <p style="text-align:center; background:#e9f2ff; padding:4px; border-radius:4px; color:#0056b3; font-weight:bold;">{role}</p>
                <p style="text-align:center; font-size:12px; margin-top:5px; color:#6c757d;">{email or user.get('Username', 'bilgi')}</p>
            </div>
            """,
            unsafe_allow_html=True
        )
        with st.sidebar:
            selected = option_menu(
                None, allowed_options, icons=icons, key="menu_corp",
                styles={
                    "container": {"background": "#f8f9fa", "border-radius": "0"},
                    "icon": {"color": "#007bff", "font-size": "18px"},
                    "nav-link": {"font-size": "16px", "color": "#212529", "border-radius": "0", "margin": "0", "--hover-color": "#e2e6ea"},
                    "nav-link-selected": {"background": "#007bff", "color": "white", "font-weight": "bold"},
                }
            )
    # =========================== YENİ TEMA 9: Retro Terminal ============================
    elif theme == "Retro Terminal":
        st.sidebar.markdown(
            """
            <style>
                .retro-card {
                    background: #1c1c1c; /* Terminal Siyahı */
                    border: 2px dashed #00ff41; /* Yeşil Işın */
                    border-radius: 0;
                    padding: 15px;
                    box-shadow: 0 0 10px rgba(0, 255, 65, 0.5); /* Yeşil Glow */
                    margin-bottom: 15px;
                    color: #00ff41; /* Yeşil Metin */
                    font-family: 'Courier New', Courier, monospace;
                    text-align: center;
                }
            </style>
            """,
            unsafe_allow_html=True
        )  
        st.sidebar.markdown(
            f"""
            <div class="retro-card">
                <img src="https://ui-avatars.com/api/?name={full_name}&background=00ff41&color=1c1c1c&size=100"
                style="border-radius:0; width:100px; display:block; margin:auto; border: 2px solid #00ff41;">
                <h3 style="text-align:center; margin-top:10px; color:#00ff41;">> {full_name}_Giriş</h3>
                <p style="text-align:center; background:rgba(0, 255, 65, 0.1); padding:4px; border-radius:0; color:#00ff41; font-weight:bold;">[{role}_YETKİLİ]</p>
                <p style="text-align:center; font-size:10px; margin-top:5px; color:#00ff41;">{email or user.get('Username', 'bilgi')}</p>
            </div>
            """,
            unsafe_allow_html=True
        )
        with st.sidebar:
            selected = option_menu(
                None, allowed_options, icons=icons, key="menu_retro",
                styles={
                    "container": {"background": "#1c1c1c", "border-radius": "0", "border": "1px solid #00ff41"},
                    "icon": {"color": "#00ff41", "font-size": "18px"},
                    "nav-link": {
                        "font-size": "16px", "color": "#00ff41", "border-radius": "0", "margin": "0", 
                        "--hover-color": "#2c2c2c", 
                        "font-family": 'Courier New'
                    },
                    "nav-link-selected": {
                        "background": "#00ff41", "color": "#1c1c1c", "font-weight": "bold",
                        "border-left": "4px solid white"
                    },
                }
            )
    # ============================ YENİ TEMA 10: Tropical Sunset ============================
    elif theme == "Tropical Sunset":
        st.sidebar.markdown(
            """
            <style>
                .tropical-card {
                    background: linear-gradient(135deg, #ff9a85 0%, #fc7e5c 100%); /* Turuncu-Pembe Gradyan */
                    border-radius: 15px;
                    padding: 25px;
                    box-shadow: 0 5px 15px rgba(255, 120, 0, 0.4); 
                    margin-bottom: 20px;
                    color: white;
                    text-shadow: 1px 1px 2px rgba(0,0,0,0.2);
                }
            </style>
            """,
            unsafe_allow_html=True
        )   
        st.sidebar.markdown(
            f"""
            <div class="tropical-card">
                <img src="https://ui-avatars.com/api/?name={full_name}&background=fec719&color=9e1c4a&size=110"
                style="border-radius:50%; width:110px; display:block; margin:auto; border: 3px solid #ffde6b;">
                <h3 style="text-align:center; margin-top:10px; color:white;">{full_name}</h3>
                <p style="text-align:center; background:rgba(255, 255, 255, 0.3); padding:4px; border-radius:10px; color:#5c102a; font-weight:bold;">{role}</p>
                <p style="text-align:center; font-size:12px; margin-top:5px; color:#ffe0b2;">{email or user.get('Username', 'bilgi')}</p>
            </div>
            """,
            unsafe_allow_html=True
        )
        with st.sidebar:
            selected = option_menu(
                None, allowed_options, icons=icons, key="menu_tropical",
                styles={
                    "container": {"background": "#fff3e0", "border-radius": "10px"},
                    "icon": {"color": "#ff7043", "font-size": "18px"},
                    "nav-link": {"font-size": "16px", "color": "#3e2723", "border-radius": "8px", "margin": "5px 0", "--hover-color": "#ffe0b2"},
                    "nav-link-selected": {
                        "background": "linear-gradient(90deg, #ffab40, #ff7043)",
                        "color": "white",
                        "font-weight": "bold"
                    },
                }
            )
    # ============================ YENİ TEMA 11: Soft Pastel ============================
    else: 
        st.sidebar.markdown(
            """
            <style>
                .pastel-card {
                    background: #fdf6f0; /* Açık Ten Rengi */
                    border: 1px solid #d4c2c2; 
                    border-radius: 12px;
                    padding: 25px;
                    box-shadow: 0 4px 6px rgba(100, 100, 100, 0.1); 
                    margin-bottom: 20px;
                    color: #5d4037; /* Kahverengi Tonu */
                    text-align: center;
                }
            </style>
            """,
            unsafe_allow_html=True
        )
        st.sidebar.markdown(
            f"""
            <div class="pastel-card">
                <img src="https://ui-avatars.com/api/?name={full_name}&background=b39ddb&color=5d4037&size=110"
                style="border-radius:50%; width:110px; display:block; margin:auto; border: 3px solid #b39ddb;">
                <h3 style="text-align:center; margin-top:10px; color:#5d4037;">{full_name}</h3>
                <p style="text-align:center; background:#f0f4c3; padding:4px; border-radius:10px; color:#558b2f; font-weight:bold;">{role}</p>
                <p style="text-align:center; font-size:12px; margin-top:5px; color:#795548;">{email or user.get('Username', 'bilgi')}</p>
            </div>
            """,
            unsafe_allow_html=True
        )
        with st.sidebar:
            selected = option_menu(
                None, allowed_options, icons=icons, key="menu_pastel",
                styles={
                    "container": {"background": "#fafafa", "border-radius": "8px"},
                    "icon": {"color": "#81c784", "font-size": "18px"},
                    "nav-link": {"font-size": "16px", "color": "#5d4037", "border-radius": "6px", "margin": "5px 0", "--hover-color": "#e8eaf6"},
                    "nav-link-selected": {
                        "background": "#b39ddb",
                        "color": "white",
                        "font-weight": "bold"
                    },
                }
            )

    st.sidebar.markdown("---") 
    if st.sidebar.button("🚪 Çıkış Yap", help="Oturumu kapatır ve uygulamayı sıfırlar."):
        for k in list(st.session_state.keys()):
            del st.session_state[k]
        st.rerun()

    return selected

# ==============================================================================
# 6. ANA AKIŞ
# ==============================================================================
def main():

    if "user" not in st.session_state:
        login_page()
        return

    user = st.session_state["user"]
    role = user["Role"]

    allowed_options = [menu for menu, roles in MENU_ROLES.items() if role in roles]
    menu_icons = {
        "Gösterge Paneli": "speedometer2",
        "Dersler": "journal-bookmark-fill", 
        "Soru Ekle": "plus-circle",
        "Soru Bankası": "collection",
        "Sınav Oluştur": "file-earmark-text",
        "Arşiv": "archive",
        "Yönetim": "gear"
    }
    icons = [menu_icons[o] for o in allowed_options]

    selected = render_sidebar(user, allowed_options, icons)

    route_page(selected)


if __name__ == "__main__":
    main()